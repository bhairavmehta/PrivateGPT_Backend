import os
import asyncio
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import logging
import json
import tempfile
import base64
from PIL import Image
import io

# Report generation imports
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from jinja2 import Environment, FileSystemLoader, Template
    HAS_JINJA2 = True
except ImportError:
    HAS_JINJA2 = False

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Service for generating reports from templates"""
    
    def __init__(self, settings):
        self.settings = settings
        self.templates_path = Path(settings.report_templates_path)
        self.output_path = Path(settings.output_path)
        
        # Ensure directories exist
        self.templates_path.mkdir(exist_ok=True)
        self.output_path.mkdir(exist_ok=True)
        
        # Initialize Jinja2 environment if available
        if HAS_JINJA2:
            self.jinja_env = Environment(loader=FileSystemLoader(str(self.templates_path)))
        else:
            self.jinja_env = None
            logger.warning("Jinja2 not available - template rendering will be limited")
        
        # Default templates will be created on first use
    
    async def generate_report(
        self,
        template_name: str,
        data: Dict[str, Any],
        output_format: str = "pdf",
        custom_filename: str = None
    ) -> str:
        """Generate a report using the specified template and data"""
        try:
            # Validate format
            if output_format.lower() not in ["pdf", "doc", "docx"]:
                raise ValueError(f"Unsupported output format: {output_format}")
            
            # Load and render template
            template_content = await self._load_template(template_name)
            rendered_content = await self._render_template(template_content, data)
            
            # Generate filename
            if custom_filename:
                filename = f"{custom_filename}.{output_format}"
            else:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"report_{template_name}_{timestamp}.{output_format}"
            
            output_file_path = self.output_path / filename
            
            # Generate report based on format
            if output_format.lower() == "pdf":
                await self._generate_pdf_report(rendered_content, data, output_file_path)
            elif output_format.lower() in ["doc", "docx"]:
                await self._generate_docx_report(rendered_content, data, output_file_path)
            
            logger.info(f"Report generated successfully: {output_file_path}")
            return str(output_file_path)
            
        except Exception as e:
            logger.error(f"Report generation failed: {e}")
            raise
    
    async def _load_template(self, template_name: str) -> Dict[str, Any]:
        """Load template configuration and content"""
        try:
            template_dir = self.templates_path / template_name
            if not template_dir.exists():
                raise FileNotFoundError(f"Template '{template_name}' not found")
            
            # Load template configuration
            config_file = template_dir / "config.json"
            if config_file.exists():
                with open(config_file, 'r') as f:
                    config = json.load(f)
            else:
                config = {"name": template_name, "description": "Auto-generated template"}
            
            # Load template content
            content_file = template_dir / "template.html"
            if content_file.exists():
                with open(content_file, 'r') as f:
                    content = f.read()
            else:
                content = self._get_default_template_content()
            
            return {
                "config": config,
                "content": content,
                "template_path": str(template_dir)
            }
            
        except Exception as e:
            logger.error(f"Failed to load template {template_name}: {e}")
            raise
    
    async def _render_template(self, template_data: Dict[str, Any], data: Dict[str, Any]) -> str:
        """Render template with provided data"""
        try:
            if self.jinja_env:
                template = self.jinja_env.from_string(template_data["content"])
                return template.render(data)
            else:
                # Simple string replacement fallback
                content = template_data["content"]
                for key, value in data.items():
                    content = content.replace(f"{{{{{key}}}}}", str(value))
                return content
                
        except Exception as e:
            logger.error(f"Template rendering failed: {e}")
            raise
    
    async def _generate_pdf_report(self, content: str, data: Dict[str, Any], output_path: Path):
        """Generate PDF report using ReportLab"""
        if not HAS_REPORTLAB:
            raise ImportError("ReportLab is required for PDF generation")
        
        try:
            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Build story (content elements)
            story = []
            styles = getSampleStyleSheet()
            
            # Parse content and convert to PDF elements
            elements = self._parse_content_for_pdf(content, data, styles)
            story.extend(elements)
            
            # Build PDF
            doc.build(story)
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise
    
    async def _generate_docx_report(self, content: str, data: Dict[str, Any], output_path: Path):
        """Generate DOCX report using python-docx"""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX generation")
        
        try:
            # Create document
            doc = Document()
            
            # Parse content and add to document
            self._parse_content_for_docx(content, data, doc)
            
            # Save document
            doc.save(str(output_path))
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise
    
    def _parse_content_for_pdf(self, content: str, data: Dict[str, Any], styles) -> List[Any]:
        """Parse HTML-like content and convert to ReportLab elements"""
        elements = []
        
        # Simple parsing - in a real implementation, you'd use a proper HTML parser
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 12))
                continue
            
            if line.startswith('<h1>') and line.endswith('</h1>'):
                text = line[4:-5]
                elements.append(Paragraph(text, styles['Title']))
                elements.append(Spacer(1, 12))
            elif line.startswith('<h2>') and line.endswith('</h2>'):
                text = line[4:-5]
                elements.append(Paragraph(text, styles['Heading1']))
                elements.append(Spacer(1, 6))
            elif line.startswith('<h3>') and line.endswith('</h3>'):
                text = line[4:-5]
                elements.append(Paragraph(text, styles['Heading2']))
                elements.append(Spacer(1, 6))
            elif line.startswith('<p>') and line.endswith('</p>'):
                text = line[3:-4]
                elements.append(Paragraph(text, styles['Normal']))
                elements.append(Spacer(1, 6))
            elif line == '<pagebreak>':
                elements.append(PageBreak())
            else:
                # Default to normal paragraph
                if line:
                    elements.append(Paragraph(line, styles['Normal']))
                    elements.append(Spacer(1, 6))
        
        return elements
    
    def _parse_content_for_docx(self, content: str, data: Dict[str, Any], doc):
        """Parse HTML-like content and add to DOCX document"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('<h1>') and line.endswith('</h1>'):
                text = line[4:-5]
                heading = doc.add_heading(text, level=1)
            elif line.startswith('<h2>') and line.endswith('</h2>'):
                text = line[4:-5]
                heading = doc.add_heading(text, level=2)
            elif line.startswith('<h3>') and line.endswith('</h3>'):
                text = line[4:-5]
                heading = doc.add_heading(text, level=3)
            elif line.startswith('<p>') and line.endswith('</p>'):
                text = line[3:-4]
                doc.add_paragraph(text)
            elif line == '<pagebreak>':
                doc.add_page_break()
            else:
                # Default to normal paragraph
                if line:
                    doc.add_paragraph(line)
    
    async def list_templates(self) -> List[Dict[str, Any]]:
        """List all available report templates"""
        try:
            # Ensure default templates exist
            await self._create_default_templates()
            
            templates = []
            
            for template_dir in self.templates_path.iterdir():
                if template_dir.is_dir():
                    config_file = template_dir / "config.json"
                    
                    if config_file.exists():
                        try:
                            with open(config_file, 'r') as f:
                                config = json.load(f)
                        except Exception as e:
                            logger.warning(f"Failed to load config for {template_dir.name}: {e}")
                            config = {"name": template_dir.name, "description": "Invalid config"}
                    else:
                        config = {"name": template_dir.name, "description": "No description available"}
                    
                    # Add template info
                    template_info = {
                        "name": template_dir.name,
                        "config": config,
                        "has_content": (template_dir / "template.html").exists(),
                        "created_at": datetime.fromtimestamp(template_dir.stat().st_ctime).isoformat(),
                        "modified_at": datetime.fromtimestamp(template_dir.stat().st_mtime).isoformat()
                    }
                    
                    templates.append(template_info)
            
            return sorted(templates, key=lambda x: x["name"])
            
        except Exception as e:
            logger.error(f"Failed to list templates: {e}")
            raise
    
    async def create_template(self, name: str, config: Dict[str, Any], content: str) -> str:
        """Create a new report template"""
        try:
            template_dir = self.templates_path / name
            template_dir.mkdir(exist_ok=True)
            
            # Save config
            config_file = template_dir / "config.json"
            with open(config_file, 'w') as f:
                json.dump(config, f, indent=2)
            
            # Save content
            content_file = template_dir / "template.html"
            with open(content_file, 'w') as f:
                f.write(content)
            
            logger.info(f"Template '{name}' created successfully")
            return str(template_dir)
            
        except Exception as e:
            logger.error(f"Failed to create template '{name}': {e}")
            raise
    
    async def delete_template(self, name: str):
        """Delete a report template"""
        try:
            template_dir = self.templates_path / name
            if template_dir.exists():
                # Remove all files in the template directory
                for file_path in template_dir.rglob("*"):
                    if file_path.is_file():
                        file_path.unlink()
                template_dir.rmdir()
                logger.info(f"Template '{name}' deleted successfully")
            else:
                raise FileNotFoundError(f"Template '{name}' not found")
                
        except Exception as e:
            logger.error(f"Failed to delete template '{name}': {e}")
            raise
    
    async def _create_default_templates(self):
        """Create default report templates if they don't exist"""
        try:
            # Research Report Template
            research_template_dir = self.templates_path / "research_report"
            if not research_template_dir.exists():
                await self.create_template(
                    "research_report",
                    {
                        "name": "Research Report",
                        "description": "Template for deep research reports",
                        "author": "Local LLM Backend",
                        "version": "1.0"
                    },
                    self._get_research_template_content()
                )
            
            # Analysis Report Template
            analysis_template_dir = self.templates_path / "analysis_report"
            if not analysis_template_dir.exists():
                await self.create_template(
                    "analysis_report",
                    {
                        "name": "Analysis Report",
                        "description": "Template for data analysis reports",
                        "author": "Local LLM Backend",
                        "version": "1.0"
                    },
                    self._get_analysis_template_content()
                )
            
            # Code Review Template
            code_review_template_dir = self.templates_path / "code_review"
            if not code_review_template_dir.exists():
                await self.create_template(
                    "code_review",
                    {
                        "name": "Code Review Report",
                        "description": "Template for code review reports",
                        "author": "Local LLM Backend",
                        "version": "1.0"
                    },
                    self._get_code_review_template_content()
                )
                
        except Exception as e:
            logger.error(f"Failed to create default templates: {e}")
    
    def _get_default_template_content(self) -> str:
        """Get default template content"""
        return """
<h1>{{title}}</h1>
<p><strong>Generated on:</strong> {{date}}</p>
<p><strong>Author:</strong> {{author}}</p>

<h2>Summary</h2>
<p>{{summary}}</p>

<h2>Content</h2>
<p>{{content}}</p>

<h2>Conclusion</h2>
<p>{{conclusion}}</p>
"""
    
    def _get_research_template_content(self) -> str:
        """Get research report template content"""
        return """
<h1>{{title}}</h1>
<p><strong>Research Topic:</strong> {{topic}}</p>
<p><strong>Generated on:</strong> {{date}}</p>
<p><strong>Research Depth:</strong> {{depth}}</p>
<p><strong>Number of Sources:</strong> {{source_count}}</p>

<h2>Executive Summary</h2>
<p>{{summary}}</p>

<h2>Research Findings</h2>
<p>{{findings}}</p>

<h2>Key Points</h2>
<p>{{key_points}}</p>

<h2>Contradictions and Uncertainties</h2>
<p>{{contradictions}}</p>

<h2>Analysis</h2>
<p>{{analysis}}</p>

<h2>Sources</h2>
<p>{{sources}}</p>

<h2>Confidence Score</h2>
<p>{{confidence_score}}</p>

<h2>Recommendations</h2>
<p>{{recommendations}}</p>
"""
    
    def _get_analysis_template_content(self) -> str:
        """Get analysis report template content"""
        return """
<h1>{{title}}</h1>
<p><strong>Analysis Type:</strong> {{analysis_type}}</p>
<p><strong>Generated on:</strong> {{date}}</p>
<p><strong>Data Source:</strong> {{data_source}}</p>

<h2>Overview</h2>
<p>{{overview}}</p>

<h2>Methodology</h2>
<p>{{methodology}}</p>

<h2>Results</h2>
<p>{{results}}</p>

<h2>Key Insights</h2>
<p>{{insights}}</p>

<h2>Recommendations</h2>
<p>{{recommendations}}</p>

<h2>Limitations</h2>
<p>{{limitations}}</p>

<h2>Appendix</h2>
<p>{{appendix}}</p>
"""
    
    def _get_code_review_template_content(self) -> str:
        """Get code review template content"""
        return """
<h1>Code Review Report</h1>
<p><strong>Project:</strong> {{project_name}}</p>
<p><strong>Reviewed on:</strong> {{date}}</p>
<p><strong>Reviewer:</strong> {{reviewer}}</p>
<p><strong>Files Reviewed:</strong> {{file_count}}</p>

<h2>Summary</h2>
<p>{{summary}}</p>

<h2>Code Quality Score</h2>
<p>{{quality_score}}</p>

<h2>Issues Found</h2>
<p>{{issues}}</p>

<h2>Recommendations</h2>
<p>{{recommendations}}</p>

<h2>Best Practices</h2>
<p>{{best_practices}}</p>

<h2>File Details</h2>
<p>{{file_details}}</p>
"""

    async def generate_vision_enhanced_report(
        self,
        template_name: str,
        data: Dict[str, Any],
        images: List[Dict[str, Any]] = None,
        output_format: str = "pdf",
        custom_filename: str = None
    ) -> str:
        """Generate a report with vision model analysis of included images"""
        try:
            # Import vision model here to avoid circular imports
            from models.llm_manager import LLMManager
            
            vision_enhanced_data = data.copy()
            
            # Process images with vision model if provided
            if images:
                logger.info(f"Processing {len(images)} images with vision model")
                llm_manager = LLMManager()
                
                image_analyses = []
                for i, image_data in enumerate(images):
                    try:
                        # Analyze image with vision model
                        analysis = await self._analyze_image_with_vision(
                            llm_manager, image_data, f"Analyze this image for report generation"
                        )
                        
                        image_analyses.append({
                            "image_index": i + 1,
                            "filename": image_data.get("filename", f"image_{i+1}"),
                            "analysis": analysis,
                            "metadata": image_data.get("metadata", {})
                        })
                        
                    except Exception as e:
                        logger.error(f"Failed to analyze image {i}: {e}")
                        image_analyses.append({
                            "image_index": i + 1,
                            "filename": image_data.get("filename", f"image_{i+1}"),
                            "analysis": f"Image analysis failed: {str(e)}",
                            "metadata": {}
                        })
                
                # Add image analyses to report data
                vision_enhanced_data["image_analyses"] = image_analyses
                vision_enhanced_data["has_images"] = True
                vision_enhanced_data["image_count"] = len(images)
                
                # Create summary of all image analyses
                all_analyses = [img["analysis"] for img in image_analyses if "failed" not in img["analysis"]]
                if all_analyses:
                    vision_enhanced_data["image_summary"] = self._create_image_summary(all_analyses)
                else:
                    vision_enhanced_data["image_summary"] = "No successful image analyses available."
            else:
                vision_enhanced_data["has_images"] = False
                vision_enhanced_data["image_count"] = 0
                vision_enhanced_data["image_analyses"] = []
                vision_enhanced_data["image_summary"] = "No images provided for analysis."
            
            # Generate report with enhanced data
            return await self.generate_report(
                template_name, vision_enhanced_data, output_format, custom_filename
            )
            
        except Exception as e:
            logger.error(f"Vision-enhanced report generation failed: {e}")
            raise
    
    async def _analyze_image_with_vision(
        self, 
        llm_manager, 
        image_data: Dict[str, Any], 
        prompt: str
    ) -> str:
        """Analyze an image using the vision model"""
        try:
            # Prepare image for vision model
            if "base64" in image_data:
                # Image is already base64 encoded
                image_base64 = image_data["base64"]
            elif "file_path" in image_data:
                # Load image from file path
                with open(image_data["file_path"], "rb") as f:
                    image_base64 = base64.b64encode(f.read()).decode('utf-8')
            elif "binary_data" in image_data:
                # Convert binary data to base64
                image_base64 = base64.b64encode(image_data["binary_data"]).decode('utf-8')
            else:
                raise ValueError("No valid image data found in image_data")
            
            # Create vision request
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image", "image": f"data:image/jpeg;base64,{image_base64}"}
                    ]
                }
            ]
            
            # Generate analysis using vision model
            response = await llm_manager.generate_response(
                messages=messages,
                model_type="vision",
                max_tokens=1024,
                temperature=0.3
            )
            
            return response.get("content", "Failed to analyze image")
            
        except Exception as e:
            logger.error(f"Image analysis failed: {e}")
            return f"Image analysis failed: {str(e)}"
    
    def _create_image_summary(self, analyses: List[str]) -> str:
        """Create a summary of multiple image analyses"""
        try:
            if not analyses:
                return "No image analyses available."
            
            if len(analyses) == 1:
                return f"Image Analysis:\n{analyses[0]}"
            
            # Create numbered summary
            summary_parts = ["Multiple Image Analyses:"]
            for i, analysis in enumerate(analyses, 1):
                summary_parts.append(f"\n{i}. {analysis}")
            
            return "\n".join(summary_parts)
            
        except Exception as e:
            logger.error(f"Failed to create image summary: {e}")
            return "Failed to create image summary."

    async def cleanup_old_reports(self, max_age_days: int = 30):
        """Clean up old generated reports"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_days * 24 * 3600
            
            for file_path in self.output_path.iterdir():
                if file_path.is_file() and file_path.suffix in ['.pdf', '.doc', '.docx']:
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        logger.info(f"Cleaned up old report: {file_path}")
                        
        except Exception as e:
            logger.error(f"Failed to cleanup old reports: {e}")
