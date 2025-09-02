import os
import mimetypes
import tempfile
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import logging
import hashlib
from datetime import datetime

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

# Document processing imports
try:
    import PyPDF2
    from docx import Document
    import markdown
    HAS_DOC_LIBS = True
except ImportError:
    HAS_DOC_LIBS = False

# Image processing imports
try:
    from PIL import Image, ExifTags
    import cv2
    import numpy as np
    HAS_IMAGE_LIBS = True
except ImportError:
    HAS_IMAGE_LIBS = False

logger = logging.getLogger(__name__)

class FileProcessor:
    """Service for processing uploaded files of various types"""
    
    def __init__(self, settings):
        self.settings = settings
        self.max_file_size = settings.max_file_size
        self.allowed_extensions = settings.allowed_file_types
        self.temp_path = Path(settings.temp_files_path)
        self.temp_path.mkdir(exist_ok=True)
        
        # File type categories
        self.image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}
        self.document_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.rtf'}
        self.code_extensions = {'.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml',
                               '.cpp', '.c', '.java', '.go', '.rs', '.php', '.rb', '.swift', '.ts', '.jsx', '.tsx'}
        
    async def process_file(self, file_content: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """Process an uploaded file and extract relevant information"""
        try:
            # Validate file
            validation_result = await self.validate_file(file_content, filename, content_type)
            if not validation_result['valid']:
                raise ValueError(validation_result['error'])
            
            file_info = validation_result['info']
            file_extension = Path(filename).suffix.lower()
            
            # Create temporary file
            temp_file_path = await self.save_temp_file(file_content, filename)
            file_info['temp_path'] = str(temp_file_path)
            
            # Process based on file type
            if file_extension in self.image_extensions:
                processing_result = await self.process_image(temp_file_path, file_info)
            elif file_extension in self.document_extensions:
                processing_result = await self.process_document(temp_file_path, file_info)
            elif file_extension in self.code_extensions:
                processing_result = await self.process_code(temp_file_path, file_info)
            else:
                processing_result = await self.process_generic_text(temp_file_path, file_info)
            
            # Merge file info with processing results
            result = {**file_info, **processing_result}
            result['processed_at'] = datetime.now().isoformat()
            
            return result
            
        except Exception as e:
            logger.error(f"File processing failed for {filename}: {e}")
            raise
    
    async def validate_file(self, file_content: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """Validate uploaded file against size and type restrictions"""
        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                return {
                    'valid': False,
                    'error': f"File size ({len(file_content)} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
                }
            
            file_extension = Path(filename).suffix.lower()
            
            # Check file extension
            if file_extension not in self.allowed_extensions:
                return {
                    'valid': False,
                    'error': f"File type '{file_extension}' is not allowed. Allowed types: {', '.join(self.allowed_extensions)}"
                }
            
            # Detect actual file type
            detected_type = await self.detect_file_type(file_content, filename)
            
            file_info = {
                'filename': filename,
                'size': len(file_content),
                'extension': file_extension,
                'mime_type': content_type,
                'detected_type': detected_type,
                'file_hash': hashlib.md5(file_content).hexdigest(),
                'category': self.categorize_file(file_extension)
            }
            
            return {'valid': True, 'info': file_info}
            
        except Exception as e:
            return {'valid': False, 'error': f"Validation failed: {str(e)}"}
    
    async def detect_file_type(self, file_content: bytes, filename: str) -> str:
        """Detect actual file type from content"""
        try:
            if HAS_MAGIC:
                mime = magic.from_buffer(file_content, mime=True)
                return mime
            else:
                # Fallback to mimetypes
                mime_type, _ = mimetypes.guess_type(filename)
                return mime_type or "application/octet-stream"
        except Exception as e:
            logger.warning(f"File type detection failed: {e}")
            return "unknown"
    
    def categorize_file(self, extension: str) -> str:
        """Categorize file by extension"""
        if extension in self.image_extensions:
            return "image"
        elif extension in self.document_extensions:
            return "document"
        elif extension in self.code_extensions:
            return "code"
        else:
            return "unknown"
    
    async def save_temp_file(self, file_content: bytes, filename: str) -> Path:
        """Save file content to temporary location"""
        try:
            # Create unique filename to avoid collisions
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_hash = hashlib.md5(file_content).hexdigest()[:8]
            safe_filename = f"{timestamp}_{file_hash}_{Path(filename).name}"
            
            temp_file_path = self.temp_path / safe_filename
            
            with open(temp_file_path, 'wb') as f:
                f.write(file_content)
            
            return temp_file_path
            
        except Exception as e:
            logger.error(f"Failed to save temp file: {e}")
            raise
    
    async def process_image(self, file_path: Path, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process image files and extract metadata"""
        try:
            result = {'type': 'image', 'analysis_ready': True}
            
            if HAS_IMAGE_LIBS:
                with Image.open(file_path) as img:
                    result.update({
                        'dimensions': img.size,
                        'mode': img.mode,
                        'format': img.format,
                        'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                    })
                    
                    # Extract EXIF data if available
                    if hasattr(img, '_getexif') and img._getexif():
                        exif = img._getexif()
                        exif_data = {}
                        for tag_id, value in exif.items():
                            tag = ExifTags.TAGS.get(tag_id, tag_id)
                            exif_data[tag] = str(value)
                        result['exif'] = exif_data
            
            return result
            
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            return {'type': 'image', 'analysis_ready': False, 'error': str(e)}
    
    async def process_document(self, file_path: Path, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process document files and extract text content"""
        try:
            result = {'type': 'document'}
            extension = file_info['extension']
            
            if extension == '.pdf':
                result.update(await self.extract_pdf_content(file_path))
            elif extension in ['.docx', '.doc']:
                result.update(await self.extract_docx_content(file_path))
            elif extension == '.txt':
                result.update(await self.extract_text_content(file_path))
            elif extension == '.md':
                result.update(await self.extract_markdown_content(file_path))
            else:
                result.update(await self.extract_text_content(file_path))
            
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {'type': 'document', 'content': '', 'error': str(e)}
    
    async def process_code(self, file_path: Path, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process code files and extract relevant information"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            result = {
                'type': 'code',
                'content': content,
                'line_count': len(content.split('\n')),
                'character_count': len(content),
                'language': self.detect_programming_language(file_info['extension'])
            }
            
            # Extract basic code metrics
            result.update(self.analyze_code_structure(content, file_info['extension']))
            
            return result
            
        except Exception as e:
            logger.error(f"Code processing failed: {e}")
            return {'type': 'code', 'content': '', 'error': str(e)}
    
    async def process_generic_text(self, file_path: Path, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process generic text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'type': 'text',
                'content': content,
                'line_count': len(content.split('\n')),
                'character_count': len(content),
                'word_count': len(content.split())
            }
            
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            return {'type': 'text', 'content': '', 'error': str(e)}
    
    async def extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text content from PDF files"""
        if not HAS_DOC_LIBS:
            return {'content': '', 'error': 'PDF processing libraries not available'}
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_content = ""
                
                for page in reader.pages:
                    text_content += page.extract_text() + "\n"
                
                return {
                    'content': text_content.strip(),
                    'page_count': len(reader.pages),
                    'has_metadata': bool(reader.metadata)
                }
                
        except Exception as e:
            return {'content': '', 'error': f"PDF extraction failed: {str(e)}"}
    
    async def extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text content from DOCX files"""
        if not HAS_DOC_LIBS:
            return {'content': '', 'error': 'DOCX processing libraries not available'}
        
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n'.join(paragraphs)
            
            return {
                'content': content,
                'paragraph_count': len(paragraphs),
                'has_tables': len(doc.tables) > 0,
                'table_count': len(doc.tables)
            }
            
        except Exception as e:
            return {'content': '', 'error': f"DOCX extraction failed: {str(e)}"}
    
    async def extract_text_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'content': content,
                'line_count': len(content.split('\n')),
                'word_count': len(content.split())
            }
            
        except Exception as e:
            return {'content': '', 'error': f"Text extraction failed: {str(e)}"}
    
    async def extract_markdown_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            result = {
                'content': content,
                'raw_markdown': content
            }
            
            if HAS_DOC_LIBS:
                # Convert to HTML if markdown library is available
                html_content = markdown.markdown(content)
                result['html_content'] = html_content
            
            return result
            
        except Exception as e:
            return {'content': '', 'error': f"Markdown extraction failed: {str(e)}"}
    
    def detect_programming_language(self, extension: str) -> str:
        """Detect programming language from file extension"""
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.jsx': 'jsx',
            '.tsx': 'tsx',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.go': 'go',
            '.rs': 'rust',
            '.php': 'php',
            '.rb': 'ruby',
            '.swift': 'swift',
            '.html': 'html',
            '.css': 'css',
            '.json': 'json',
            '.xml': 'xml',
            '.yaml': 'yaml',
            '.yml': 'yaml'
        }
        return language_map.get(extension, 'unknown')
    
    def analyze_code_structure(self, content: str, extension: str) -> Dict[str, Any]:
        """Analyze basic code structure"""
        lines = content.split('\n')
        
        # Basic metrics
        blank_lines = sum(1 for line in lines if not line.strip())
        comment_lines = 0
        
        # Language-specific comment detection
        if extension in ['.py']:
            comment_lines = sum(1 for line in lines if line.strip().startswith('#'))
        elif extension in ['.js', '.ts', '.jsx', '.tsx', '.java', '.cpp', '.c', '.go', '.php', '.swift']:
            comment_lines = sum(1 for line in lines if line.strip().startswith('//'))
        elif extension in ['.html', '.xml']:
            comment_lines = sum(1 for line in lines if '<!--' in line)
        elif extension in ['.css']:
            comment_lines = sum(1 for line in lines if '/*' in line)
        
        return {
            'blank_lines': blank_lines,
            'comment_lines': comment_lines,
            'code_lines': len(lines) - blank_lines - comment_lines,
            'code_quality_score': self.calculate_code_quality_score(lines, extension)
        }
    
    def calculate_code_quality_score(self, lines: List[str], extension: str) -> float:
        """Calculate a basic code quality score"""
        if not lines:
            return 0.0
        
        total_lines = len(lines)
        blank_lines = sum(1 for line in lines if not line.strip())
        long_lines = sum(1 for line in lines if len(line) > 120)
        
        # Basic scoring (this is a simplified example)
        score = 1.0
        
        # Penalize files with no blank lines (likely minified)
        if blank_lines == 0 and total_lines > 10:
            score -= 0.3
        
        # Penalize files with too many long lines
        if long_lines > total_lines * 0.5:
            score -= 0.2
        
        return max(0.0, min(1.0, score))
    
    async def cleanup_temp_file(self, file_path: str):
        """Clean up temporary file"""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
                logger.info(f"Cleaned up temp file: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to cleanup temp file {file_path}: {e}")
    
    async def cleanup_old_temp_files(self, max_age_hours: int = 24):
        """Clean up old temporary files"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for file_path in self.temp_path.iterdir():
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        logger.info(f"Cleaned up old temp file: {file_path}")
                        
        except Exception as e:
            logger.error(f"Failed to cleanup old temp files: {e}")
