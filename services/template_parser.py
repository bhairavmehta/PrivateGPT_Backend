import os
import re
import json
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
import logging
from datetime import datetime

# Document processing imports
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

logger = logging.getLogger(__name__)

@dataclass
class TemplateSection:
    """Represents a section in a template with its structure and formatting."""
    level: int  # Heading level (1=main, 2=sub, etc.)
    title: str
    content_type: str  # 'heading', 'paragraph', 'list', 'table', 'chart_placeholder'
    formatting: Dict[str, Any] = field(default_factory=dict)
    content_pattern: str = ""  # Pattern for content generation
    subsections: List['TemplateSection'] = field(default_factory=list)
    order: int = 0

@dataclass
class TemplateStructure:
    """Complete template structure with metadata."""
    name: str
    title_pattern: str
    sections: List[TemplateSection] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    formatting_rules: Dict[str, Any] = field(default_factory=dict)
    content_guidelines: List[str] = field(default_factory=list)

class TemplateParser:
    """Advanced template parser for extracting report structures from uploaded documents."""
    
    def __init__(self, settings=None):
        self.settings = settings
        self.templates_dir = Path("templates/uploaded") if not settings else Path(settings.temp_files_path) / "templates"
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Common heading patterns
        self.heading_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings
            r'^(\d+\.?\s*.+)$',  # Numbered headings
            r'^([A-Z][A-Z\s]+)$',  # ALL CAPS headings
            r'^([A-Z][a-z\s]+):?\s*$',  # Title Case headings
            r'^(Executive Summary|Introduction|Conclusion|Recommendations?|Analysis|Findings|Overview|Background|Methodology|Results|Discussion|References?)$',  # Common section names
        ]
        
        # Content type patterns
        self.list_patterns = [
            r'^\s*[-•*]\s+(.+)$',  # Bullet points
            r'^\s*\d+\.\s+(.+)$',  # Numbered lists
            r'^\s*[a-zA-Z]\.\s+(.+)$',  # Lettered lists
        ]
    
    async def parse_template_file(self, file_path: str, template_name: str) -> TemplateStructure:
        """Parse an uploaded template file and extract its structure."""
        try:
            file_path = Path(file_path)
            file_ext = file_path.suffix.lower()
            
            logger.info(f"Parsing template file: {file_path} (type: {file_ext})")
            
            # Extract text content based on file type
            if file_ext == '.pdf' and HAS_PDF:
                content = self._extract_pdf_content(file_path)
            elif file_ext in ['.docx', '.doc'] and HAS_DOCX:
                content = self._extract_docx_content(file_path)
            elif file_ext == '.txt':
                content = self._extract_txt_content(file_path)
            elif file_ext == '.md':
                content = self._extract_markdown_content(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Analyze and extract structure
            structure = self._analyze_document_structure(content, template_name)
            
            # Save template structure
            await self._save_template_structure(structure, template_name)
            
            logger.info(f"Successfully parsed template '{template_name}' with {len(structure.sections)} sections")
            return structure
            
        except Exception as e:
            logger.error(f"Failed to parse template file {file_path}: {e}")
            raise
    
    def _extract_pdf_content(self, file_path: Path) -> List[str]:
        """Extract text content from PDF file."""
        lines = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        lines.extend(text.split('\n'))
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            raise
        return [line.strip() for line in lines if line.strip()]
    
    def _extract_docx_content(self, file_path: Path) -> List[str]:
        """Extract text content from DOCX file with formatting information."""
        lines = []
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Add formatting metadata
                    style = paragraph.style.name if paragraph.style else 'Normal'
                    text = paragraph.text.strip()
                    
                    # Mark headings
                    if 'Heading' in style:
                        level = int(style.replace('Heading ', '')) if style.replace('Heading ', '').isdigit() else 1
                        text = f"[H{level}] {text}"
                    
                    lines.append(text)
                    
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        lines.append(f"[TABLE] {row_text}")
                        
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            raise
        return lines
    
    def _extract_txt_content(self, file_path: Path) -> List[str]:
        """Extract content from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return [line.strip() for line in content.split('\n') if line.strip()]
        except Exception as e:
            logger.error(f"Error extracting TXT content: {e}")
            raise
    
    def _extract_markdown_content(self, file_path: Path) -> List[str]:
        """Extract content from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return [line.strip() for line in content.split('\n') if line.strip()]
        except Exception as e:
            logger.error(f"Error extracting Markdown content: {e}")
            raise
    
    def _analyze_document_structure(self, lines: List[str], template_name: str) -> TemplateStructure:
        """Analyze document lines and extract structural patterns."""
        sections = []
        current_section = None
        section_order = 0
        
        # Detect title pattern
        title_pattern = self._detect_title_pattern(lines)
        
        # Analyze each line
        for i, line in enumerate(lines):
            if not line.strip():
                continue
                
            # Check if line is a heading
            heading_info = self._detect_heading(line, i, lines)
            if heading_info:
                level, title = heading_info
                
                # Save previous section if exists
                if current_section:
                    sections.append(current_section)
                
                # Create new section
                current_section = TemplateSection(
                    level=level,
                    title=title,
                    content_type='heading',
                    order=section_order
                )
                section_order += 1
                
            elif current_section:
                # Analyze content type
                content_type = self._detect_content_type(line)
                
                # Add content pattern to current section
                if not current_section.content_pattern:
                    current_section.content_pattern = content_type
                elif content_type != current_section.content_pattern:
                    current_section.content_pattern = 'mixed'
        
        # Add final section
        if current_section:
            sections.append(current_section)
        
        # Extract formatting rules and guidelines
        formatting_rules = self._extract_formatting_rules(lines)
        content_guidelines = self._generate_content_guidelines(sections)
        
        return TemplateStructure(
            name=template_name,
            title_pattern=title_pattern,
            sections=sections,
            metadata={
                'created_at': datetime.now().isoformat(),
                'total_sections': len(sections),
                'max_heading_level': max([s.level for s in sections] + [1]),
                'content_types': list(set([s.content_pattern for s in sections]))
            },
            formatting_rules=formatting_rules,
            content_guidelines=content_guidelines
        )
    
    def _detect_title_pattern(self, lines: List[str]) -> str:
        """Detect the title pattern from the document."""
        if not lines:
            return "Professional Report Title"
        
        # Usually the first significant line is the title
        for line in lines[:5]:  # Check first 5 lines
            if len(line) > 10 and not any(pattern in line.lower() for pattern in ['page', 'date', 'author']):
                return line
        
        return lines[0] if lines else "Professional Report Title"
    
    def _detect_heading(self, line: str, index: int, lines: List[str]) -> Optional[Tuple[int, str]]:
        """Detect if a line is a heading and determine its level."""
        
        # Check for explicit heading markers (from DOCX)
        if line.startswith('[H'):
            match = re.match(r'\[H(\d+)\]\s*(.+)', line)
            if match:
                level = int(match.group(1))
                title = match.group(2).strip()
                return (level, title)
        
        # Check markdown headings
        markdown_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if markdown_match:
            level = len(markdown_match.group(1))
            title = markdown_match.group(2).strip()
            return (level, title)
        
        # Check numbered headings
        numbered_match = re.match(r'^(\d+\.?\d*\.?\d*)\s+(.+)$', line)
        if numbered_match and len(line) < 100:  # Likely a heading if short
            level = len(numbered_match.group(1).split('.'))
            title = numbered_match.group(2).strip()
            return (level, title)
        
        # Check ALL CAPS (likely main headings)
        if line.isupper() and 10 <= len(line) <= 80:
            return (1, line.strip())
        
        # Check common section names
        common_sections = [
            'executive summary', 'introduction', 'background', 'overview',
            'methodology', 'analysis', 'findings', 'results', 'discussion',
            'conclusion', 'conclusions', 'recommendations', 'references',
            'appendix', 'abstract', 'summary'
        ]
        
        if any(section in line.lower() for section in common_sections) and len(line) < 50:
            return (1, line.strip())
        
        # Check if line is significantly shorter than surrounding lines (might be heading)
        if index > 0 and index < len(lines) - 1:
            prev_len = len(lines[index - 1])
            next_len = len(lines[index + 1])
            current_len = len(line)
            
            if current_len < 50 and current_len < prev_len * 0.6 and current_len < next_len * 0.6:
                # Check if it's title case or has special formatting
                if line.istitle() or ':' in line:
                    return (2, line.strip())
        
        return None
    
    def _detect_content_type(self, line: str) -> str:
        """Detect the type of content in a line."""
        
        # Check for table content
        if line.startswith('[TABLE]'):
            return 'table'
        
        # Check for lists
        for pattern in self.list_patterns:
            if re.match(pattern, line):
                return 'list'
        
        # Check for chart/figure references
        if any(keyword in line.lower() for keyword in ['figure', 'chart', 'graph', 'table', 'diagram']):
            return 'chart_reference'
        
        # Default to paragraph
        return 'paragraph'
    
    def _extract_formatting_rules(self, lines: List[str]) -> Dict[str, Any]:
        """Extract formatting rules from the document."""
        rules = {
            'heading_style': 'numbered',  # Default
            'list_style': 'bullet',
            'paragraph_spacing': 'normal',
            'numbering_pattern': 'decimal'
        }
        
        # Analyze heading patterns
        numbered_headings = sum(1 for line in lines if re.match(r'^\d+\.', line))
        if numbered_headings > 2:
            rules['heading_style'] = 'numbered'
        
        # Analyze list patterns
        bullet_lists = sum(1 for line in lines if re.match(r'^\s*[-•*]\s+', line))
        numbered_lists = sum(1 for line in lines if re.match(r'^\s*\d+\.\s+', line))
        
        if bullet_lists > numbered_lists:
            rules['list_style'] = 'bullet'
        else:
            rules['list_style'] = 'numbered'
        
        return rules
    
    def _generate_content_guidelines(self, sections: List[TemplateSection]) -> List[str]:
        """Generate content guidelines based on the template structure."""
        guidelines = []
        
        # Analyze section patterns
        if any(s.title.lower() == 'executive summary' for s in sections):
            guidelines.append("Include an Executive Summary section at the beginning")
        
        if any(s.title.lower() in ['conclusion', 'conclusions'] for s in sections):
            guidelines.append("End with a Conclusion section")
        
        if any(s.title.lower() in ['recommendation', 'recommendations'] for s in sections):
            guidelines.append("Include actionable recommendations")
        
        # Analyze content types
        content_types = [s.content_pattern for s in sections]
        if 'list' in content_types:
            guidelines.append("Use bullet points or numbered lists for key information")
        
        if 'table' in content_types:
            guidelines.append("Include tables for data presentation where appropriate")
        
        guidelines.append(f"Structure the report with {len(sections)} main sections")
        
        return guidelines
    
    async def _save_template_structure(self, structure: TemplateStructure, template_name: str):
        """Save the extracted template structure to a JSON file."""
        try:
            template_file = self.templates_dir / f"{template_name}.json"
            
            # Convert dataclasses to dict for JSON serialization
            structure_dict = {
                'name': structure.name,
                'title_pattern': structure.title_pattern,
                'sections': [
                    {
                        'level': s.level,
                        'title': s.title,
                        'content_type': s.content_type,
                        'formatting': s.formatting,
                        'content_pattern': s.content_pattern,
                        'order': s.order
                    }
                    for s in structure.sections
                ],
                'metadata': structure.metadata,
                'formatting_rules': structure.formatting_rules,
                'content_guidelines': structure.content_guidelines
            }
            
            with open(template_file, 'w', encoding='utf-8') as f:
                json.dump(structure_dict, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Template structure saved to {template_file}")
            
        except Exception as e:
            logger.error(f"Failed to save template structure: {e}")
            raise
    
    async def load_template_structure(self, template_name: str) -> Optional[TemplateStructure]:
        """Load a saved template structure."""
        try:
            template_file = self.templates_dir / f"{template_name}.json"
            
            if not template_file.exists():
                return None
            
            with open(template_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert dict back to dataclass
            sections = [
                TemplateSection(
                    level=s['level'],
                    title=s['title'],
                    content_type=s['content_type'],
                    formatting=s.get('formatting', {}),
                    content_pattern=s.get('content_pattern', ''),
                    order=s.get('order', 0)
                )
                for s in data['sections']
            ]
            
            return TemplateStructure(
                name=data['name'],
                title_pattern=data['title_pattern'],
                sections=sections,
                metadata=data.get('metadata', {}),
                formatting_rules=data.get('formatting_rules', {}),
                content_guidelines=data.get('content_guidelines', [])
            )
            
        except Exception as e:
            logger.error(f"Failed to load template structure: {e}")
            return None
    
    async def list_available_templates(self) -> List[Dict[str, Any]]:
        """List all available templates with their metadata."""
        templates = []
        
        try:
            for template_file in self.templates_dir.glob("*.json"):
                structure = await self.load_template_structure(template_file.stem)
                if structure:
                    templates.append({
                        'name': structure.name,
                        'file_name': template_file.stem,
                        'sections_count': len(structure.sections),
                        'created_at': structure.metadata.get('created_at'),
                        'content_types': structure.metadata.get('content_types', []),
                        'guidelines': structure.content_guidelines[:3]  # First 3 guidelines
                    })
        except Exception as e:
            logger.error(f"Failed to list templates: {e}")
        
        return templates
    
    async def delete_template(self, template_name: str) -> bool:
        """Delete a template."""
        try:
            template_file = self.templates_dir / f"{template_name}.json"
            if template_file.exists():
                template_file.unlink()
                logger.info(f"Deleted template: {template_name}")
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to delete template: {e}")
            return False


def get_template_parser(settings=None) -> TemplateParser:
    """Get a template parser instance."""
    return TemplateParser(settings) 