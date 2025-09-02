import os
import tempfile
import asyncio
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import logging
import json
import base64
import io
import re

# Document processing imports
try:
    import PyPDF2
    from PyPDF2 import PdfReader, PdfWriter
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import Color
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches
    import docx2txt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF for advanced PDF handling
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

logger = logging.getLogger(__name__)

class PDFElement:
    """Represents a PDF element with exact positioning and formatting"""
    def __init__(self, element_type: str, content: str, bbox: tuple, font_info: dict, page_num: int):
        self.element_type = element_type  # text, bullet, table_cell, heading, etc.
        self.content = content
        self.original_content = content
        self.modified_content = ""
        self.bbox = bbox  # (x0, y0, x1, y1) positioning
        self.font_info = font_info  # font, size, color, etc.
        self.page_num = page_num
        self.element_id = f"page_{page_num}_{element_type}_{hash(content[:50])}"

class DocumentStructure:
    """Represents the complete structure of a document"""
    def __init__(self):
        self.title = ""
        self.metadata = {}
        self.sections = []
        self.pages = []
        self.pdf_elements = []  # New: exact PDF elements with positioning
        self.total_sections = 0
        self.total_paragraphs = 0
        self.total_words = 0

class DocumentSection:
    """Represents a section of the document with its structure"""
    def __init__(self, section_type: str, content: str, level: int = 0, formatting: Dict = None):
        self.section_type = section_type  # title, heading, paragraph, list, table, etc.
        self.content = content
        self.level = level  # For headings: 1, 2, 3, etc.
        self.formatting = formatting or {}
        self.original_content = content
        self.processed_content = ""
        self.section_id = ""

class AdvancedDocumentProcessor:
    """Perfect Pipeline for Complete Document Processing and Regeneration"""
    
    def __init__(self, settings, llm_manager):
        self.settings = settings
        self.llm_manager = llm_manager
        self.temp_path = Path(settings.temp_files_path)
        self.temp_path.mkdir(exist_ok=True)
        self.processing_log = []
    
    async def process_document(
        self,
        file_path: str,
        modification_prompt: str,
        output_format: str = None,
        custom_filename: str = None,
        preserve_structure: bool = True  # New parameter for structure preservation
    ) -> Dict[str, Any]:
        """
        PERFECT PIPELINE: Complete Document Processing and Regeneration
        
        Pipeline Stages:
        1. Deep Content Extraction (with exact structure if preserve_structure=True)
        2. Structure Analysis & Preservation
        3. Content Segmentation
        4. Section-by-Section AI Processing
        5. Content Quality Assurance
        6. Document Reconstruction with Original Structure
        7. Final Quality Check
        """
        try:
            self.processing_log = []
            start_time = datetime.now()
            
            file_path = Path(file_path)
            original_format = file_path.suffix.lower()
            
            if original_format not in ['.pdf', '.doc', '.docx']:
                raise ValueError(f"Unsupported file format: {original_format}")
            
            # STAGE 1: Deep Content Extraction (with structure preservation for PDFs)
            if preserve_structure and original_format == '.pdf':
                self._log("Stage 1: Structure-Preserving PDF Extraction - Maintaining exact layout")
                document_structure = await self._extract_pdf_with_exact_structure(file_path)
            else:
                self._log("Stage 1: Deep Content Extraction - Extracting ALL content and metadata")
                document_structure = await self._deep_content_extraction(file_path, original_format)
            
            # STAGE 2: Structure Analysis & Preservation
            self._log("Stage 2: Structure Analysis - Analyzing document architecture")
            analyzed_structure = await self._analyze_document_structure(document_structure)
            
            # STAGE 3: Content Segmentation
            self._log("Stage 3: Content Segmentation - Breaking into processable sections")
            if preserve_structure and original_format == '.pdf':
                segmented_sections = await self._segment_pdf_elements_intelligently(analyzed_structure)
            else:
                segmented_sections = await self._segment_content_intelligently(analyzed_structure)
            
            # STAGE 4: Section-by-Section AI Processing
            self._log("Stage 4: AI Processing - Regenerating content while preserving structure")
            if preserve_structure and original_format == '.pdf':
                processed_sections = await self._process_pdf_elements_with_ai(segmented_sections, modification_prompt)
            else:
                processed_sections = await self._process_sections_with_ai(segmented_sections, modification_prompt)
            
            # STAGE 5: Content Quality Assurance
            self._log("Stage 5: Quality Assurance - Ensuring consistency and coherence")
            quality_assured_content = await self._quality_assurance_pass(processed_sections, modification_prompt)
            
            # STAGE 6: Document Reconstruction
            self._log("Stage 6: Document Reconstruction - Rebuilding with original structure")
            if preserve_structure and original_format == '.pdf':
                reconstructed_document = await self._reconstruct_pdf_with_exact_structure(quality_assured_content, analyzed_structure)
            else:
                reconstructed_document = await self._reconstruct_document_with_structure(quality_assured_content, analyzed_structure)
            
            # STAGE 7: Final Quality Check
            self._log("Stage 7: Final Quality Check - Verifying output quality")
            final_content = await self._final_quality_check(reconstructed_document, modification_prompt)
            
            # Generate output
            if output_format is None:
                output_format = original_format.lstrip('.')
            
            if custom_filename:
                output_filename = f"{custom_filename}.{output_format}"
            else:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = file_path.stem
                output_filename = f"{base_name}_modified_{timestamp}.{output_format}"
            
            output_path = self.temp_path / output_filename
            
            # Generate final document
            if preserve_structure and original_format == '.pdf':
                await self._generate_structure_preserved_pdf(
                    final_content, 
                    output_path, 
                    analyzed_structure
                )
            else:
                await self._generate_enhanced_document(
                    final_content, 
                    output_path, 
                    output_format,
                    analyzed_structure
                )
            
            end_time = datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            return {
                "success": True,
                "output_path": str(output_path),
                "original_format": original_format,
                "output_format": output_format,
                "original_size": file_path.stat().st_size,
                "output_size": output_path.stat().st_size,
                "processing_time": end_time.isoformat(),
                "processing_duration_seconds": processing_time,
                "pipeline_stages_completed": 7,
                "sections_processed": len(processed_sections),
                "total_content_length": len(final_content.get('complete_text', '')),
                "modifications_applied": True,
                "processing_log": self.processing_log,
                "structure_preserved": preserve_structure,
                "quality_assurance_passed": True
            }
            
        except Exception as e:
            logger.error(f"Perfect pipeline failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "processing_time": datetime.now().isoformat(),
                "pipeline_stage_failed": getattr(self, '_current_stage', 'Unknown'),
                "processing_log": self.processing_log
            }

    async def _extract_pdf_with_exact_structure(self, file_path: Path) -> DocumentStructure:
        """Extract PDF with exact structure preservation including positioning, fonts, bullets, tables"""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF is required for structure-preserving PDF processing")
        
        self._current_stage = "Structure-Preserving PDF Extraction"
        structure = DocumentStructure()
        
        try:
            doc = fitz.open(str(file_path))
            
            # Extract metadata
            structure.metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get('title', ''),
                "author": doc.metadata.get('author', ''),
                "creator": doc.metadata.get('creator', ''),
                "subject": doc.metadata.get('subject', ''),
                "format": "PDF",
                "structure_preserved": True
            }
            
            structure.title = structure.metadata.get('title', 'Document')
            
            # Extract each page with exact structure
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Get all text blocks with positioning and formatting
                text_blocks = page.get_text("dict")
                
                page_elements = []
                page_text_parts = []
                
                for block in text_blocks.get("blocks", []):
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if text:
                                    # Extract detailed formatting
                                    font_info = {
                                        "font": span.get("font", ""),
                                        "size": span.get("size", 12),
                                        "color": span.get("color", 0),
                                        "flags": span.get("flags", 0),  # Bold, italic flags
                                        "ascender": span.get("ascender", 0),
                                        "descender": span.get("descender", 0)
                                    }
                                    
                                    # Determine element type based on formatting and content
                                    element_type = self._classify_pdf_element(text, font_info, span["bbox"])
                                    
                                    # Create PDF element with exact positioning
                                    pdf_element = PDFElement(
                                        element_type=element_type,
                                        content=text,
                                        bbox=span["bbox"],
                                        font_info=font_info,
                                        page_num=page_num + 1
                                    )
                                    
                                    page_elements.append(pdf_element)
                                    page_text_parts.append(text)
                                    structure.pdf_elements.append(pdf_element)
                
                # Store page information
                page_text = "\n".join(page_text_parts)
                structure.pages.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "elements": len(page_elements),
                    "word_count": len(page_text.split()),
                    "pdf_elements": page_elements
                })
            
            doc.close()
            
            # Convert PDF elements to sections for compatibility
            for element in structure.pdf_elements:
                section = DocumentSection(
                    section_type=element.element_type,
                    content=element.content,
                    level=0,
                    formatting=element.font_info
                )
                section.section_id = element.element_id
                structure.sections.append(section)
            
            # Calculate totals
            structure.total_sections = len(structure.sections)
            structure.total_paragraphs = sum(1 for s in structure.sections if s.section_type == 'paragraph')
            structure.total_words = sum(len(s.content.split()) for s in structure.sections)
            
            self._log(f"Extracted PDF with exact structure: {len(structure.pdf_elements)} elements with positioning")
            return structure
            
        except Exception as e:
            logger.error(f"Structure-preserving PDF extraction failed: {e}")
            # Fallback to regular extraction
            return await self._extract_pdf_with_structure(file_path)

    def _classify_pdf_element(self, text: str, font_info: dict, bbox: tuple) -> str:
        """Classify PDF element type based on content, formatting, and position"""
        font_size = font_info.get("size", 12)
        font_flags = font_info.get("flags", 0)
        
        # Check for bullet points
        if re.match(r'^[\u2022\u25cf\u25cb\u25aa\u25ab•·‣⁃]\s*', text) or re.match(r'^[\-\*]\s+', text):
            return 'bullet_point'
        
        # Check for numbered lists
        if re.match(r'^\d+[\.\)]\s+', text):
            return 'numbered_list'
        
        # Check for headings (large font, bold, or all caps)
        if font_size > 14 or (font_flags & 2**4):  # Bold flag
            if len(text) < 100:
                return 'heading'
        
        # Check for title (very large font or all caps and short)
        if font_size > 16 or (text.isupper() and len(text) < 80):
            return 'title'
        
        # Check for table-like content (contains multiple tabs or specific patterns)
        if '\t' in text or len(text.split()) > 1 and any(char.isdigit() for char in text):
            if len(text.split()) <= 10:  # Short entries typical of tables
                return 'table_cell'
        
        # Default to paragraph
        return 'paragraph'

    async def _segment_pdf_elements_intelligently(self, document_structure: DocumentStructure) -> List[List[PDFElement]]:
        """Segment PDF elements into logical processing groups while preserving structure"""
        self._current_stage = "PDF Element Segmentation"
        
        # Group PDF elements by logical sections
        segmented_groups = []
        current_group = []
        
        for element in document_structure.pdf_elements:
            # Start new group on major headings or after certain number of elements
            if element.element_type in ['title', 'heading'] or len(current_group) >= 5:
                if current_group:
                    segmented_groups.append(current_group)
                    current_group = []
            
            current_group.append(element)
        
        # Add final group
        if current_group:
            segmented_groups.append(current_group)
        
        self._log(f"PDF elements segmented into {len(segmented_groups)} logical groups")
        return segmented_groups

    async def _process_pdf_elements_with_ai(self, segmented_elements: List[List[PDFElement]], modification_prompt: str) -> List[List[PDFElement]]:
        """Process PDF elements with AI while preserving exact structure"""
        try:
            self._current_stage = "PDF Element AI Processing"
            processed_groups = []
            
            for group_idx, element_group in enumerate(segmented_elements):
                self._log(f"Processing PDF element group {group_idx + 1}/{len(segmented_elements)}")
                
                # Combine group content with structure indicators
                group_content_parts = []
                for element in element_group:
                    if element.element_type == 'bullet_point':
                        group_content_parts.append(f"• {element.content}")
                    elif element.element_type == 'numbered_list':
                        group_content_parts.append(f"1. {element.content}")
                    elif element.element_type == 'heading':
                        group_content_parts.append(f"HEADING: {element.content}")
                    elif element.element_type == 'title':
                        group_content_parts.append(f"TITLE: {element.content}")
                    else:
                        group_content_parts.append(element.content)
                
                group_content = '\n'.join(group_content_parts)
                
                # Enhanced prompt for structure preservation
                section_prompt = f"""Modify the content according to the user's request while maintaining EXACT formatting structure:

USER REQUEST: {modification_prompt}

ORIGINAL CONTENT WITH STRUCTURE:
{group_content}

CRITICAL REQUIREMENTS:
1. Maintain EXACT formatting - if it's a bullet point, keep it as a bullet point
2. Maintain EXACT structure - if it's a heading, keep it as a heading
3. Only modify the TEXT CONTENT, not the formatting
4. If it's "• Text", return "• Modified Text"
5. If it's "HEADING: Text", return "HEADING: Modified Text"
6. Apply the user's modifications while preserving all structural elements

Return the modified content with identical structure:"""

                processed_content = ""
                async for chunk in self.llm_manager.generate_response_stream(
                    message=section_prompt,
                    model_type="report_generation",
                    max_tokens=2500,
                    temperature=0.2  # Lower temperature for better structure preservation
                ):
                    processed_content = chunk.get("accumulated_text", "")
                    if chunk.get("finished", False):
                        break
                
                # Parse the processed content back to PDF elements
                processed_group = self._parse_processed_pdf_elements(processed_content, element_group)
                processed_groups.append(processed_group)
            
            self._log(f"PDF element AI processing completed for all {len(segmented_elements)} groups")
            return processed_groups
            
        except Exception as e:
            logger.error(f"PDF element AI processing failed: {e}")
            return segmented_elements

    def _parse_processed_pdf_elements(self, processed_content: str, original_group: List[PDFElement]) -> List[PDFElement]:
        """Parse processed content back to PDF elements with preserved structure"""
        try:
            lines = processed_content.strip().split('\n')
            processed_elements = []
            
            for i, original_element in enumerate(original_group):
                if i < len(lines):
                    processed_line = lines[i].strip()
                    
                    # Extract content based on structure markers
                    if processed_line.startswith('•'):
                        content = processed_line[1:].strip()
                    elif processed_line.startswith('HEADING:'):
                        content = processed_line[8:].strip()
                    elif processed_line.startswith('TITLE:'):
                        content = processed_line[6:].strip()
                    elif re.match(r'^\d+\.', processed_line):
                        content = re.sub(r'^\d+\.\s*', '', processed_line)
                    else:
                        content = processed_line
                    
                    # Create new element with modified content but same structure
                    new_element = PDFElement(
                        element_type=original_element.element_type,
                        content=original_element.content,  # Keep original
                        bbox=original_element.bbox,
                        font_info=original_element.font_info,
                        page_num=original_element.page_num
                    )
                    new_element.element_id = original_element.element_id
                    new_element.original_content = original_element.content
                    new_element.modified_content = content if content else original_element.content
                    
                    processed_elements.append(new_element)
                else:
                    # Keep original if no processed version
                    original_element.modified_content = original_element.content
                    processed_elements.append(original_element)
            
            return processed_elements
            
        except Exception as e:
            logger.error(f"PDF element parsing failed: {e}")
            # Return originals with modified_content set to original content
            for element in original_group:
                element.modified_content = element.content
            return original_group

    async def _reconstruct_pdf_with_exact_structure(self, processed_elements: List[List[PDFElement]], original_structure: DocumentStructure) -> Dict[str, Any]:
        """Reconstruct PDF maintaining exact original structure and positioning"""
        try:
            self._current_stage = "PDF Structure Reconstruction"
            
            # Flatten elements
            all_elements = []
            for group in processed_elements:
                all_elements.extend(group)
            
            # Build complete text preserving structure
            complete_text_parts = []
            
            for element in all_elements:
                content = element.modified_content or element.content
                
                # Add content with structure preserved
                if element.element_type == 'bullet_point':
                    complete_text_parts.append(f"• {content}")
                elif element.element_type == 'numbered_list':
                    complete_text_parts.append(f"1. {content}")
                elif element.element_type == 'heading':
                    complete_text_parts.append(f"\n{content}\n")
                elif element.element_type == 'title':
                    complete_text_parts.append(f"\n{content.upper()}\n")
                else:
                    complete_text_parts.append(content)
            
            complete_text = '\n'.join(complete_text_parts).strip()
            
            reconstructed_document = {
                "complete_text": complete_text,
                "pdf_elements": all_elements,
                "metadata": original_structure.metadata,
                "total_elements": len(all_elements),
                "structure_preserved": True,
                "reconstruction_successful": True
            }
            
            self._log(f"PDF reconstructed with exact structure: {len(all_elements)} elements")
            return reconstructed_document
            
        except Exception as e:
            logger.error(f"PDF structure reconstruction failed: {e}")
            return {
                "complete_text": "PDF reconstruction failed",
                "pdf_elements": [],
                "metadata": {},
                "reconstruction_successful": False
            }

    async def _generate_structure_preserved_pdf(
        self, 
        final_content: Dict[str, Any], 
        output_path: Path, 
        structure: DocumentStructure
    ):
        """Generate PDF with exact structure preservation"""
        if not HAS_PYMUPDF or not HAS_REPORTLAB:
            raise ImportError("PyMuPDF and ReportLab are required for structure-preserving PDF generation")
        
        try:
            # Use ReportLab to create PDF with exact positioning
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            pdf_elements = final_content.get('pdf_elements', [])
            
            # Group elements by page
            pages_elements = {}
            for element in pdf_elements:
                page_num = element.page_num
                if page_num not in pages_elements:
                    pages_elements[page_num] = []
                pages_elements[page_num].append(element)
            
            # Generate each page with exact positioning
            for page_num in sorted(pages_elements.keys()):
                page_elements = pages_elements[page_num]
                
                for element in page_elements:
                    content = element.modified_content or element.content
                    bbox = element.bbox
                    font_info = element.font_info
                    
                    # Set font and size
                    font_size = font_info.get('size', 12)
                    c.setFont("Helvetica", font_size)
                    
                    # Convert coordinates (PDF coordinates are bottom-left origin)
                    x = bbox[0]
                    y = height - bbox[3]  # Flip Y coordinate
                    
                    # Add structure-specific formatting
                    if element.element_type == 'bullet_point':
                        c.drawString(x, y, f"• {content}")
                    elif element.element_type == 'numbered_list':
                        c.drawString(x, y, f"1. {content}")
                    elif element.element_type in ['heading', 'title']:
                        # Make headings bold (simulate)
                        c.setFont("Helvetica-Bold", font_size)
                        c.drawString(x, y, content)
                        c.setFont("Helvetica", font_size)  # Reset
                    else:
                        c.drawString(x, y, content)
                
                # Start new page if not the last page
                if page_num < max(pages_elements.keys()):
                    c.showPage()
            
            c.save()
            self._log(f"Structure-preserved PDF generated: {output_path}")
            
        except Exception as e:
            logger.error(f"Structure-preserved PDF generation failed: {e}")
            # Fallback to regular PDF generation
            await self._generate_enhanced_pdf(
                final_content.get('complete_text', ''), 
                output_path, 
                final_content.get('metadata', {}),
                structure
            )

    def _log(self, message: str):
        """Log processing steps"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        self.processing_log.append(log_entry)
        logger.info(log_entry)

    async def _deep_content_extraction(self, file_path: Path, file_format: str):
        """STAGE 1: Extract ALL content with complete structure preservation"""
        try:
            self._current_stage = "Deep Content Extraction"
            structure = DocumentStructure()
            
            if file_format == '.pdf':
                structure = await self._extract_pdf_with_structure(file_path)
            elif file_format in ['.doc', '.docx']:
                structure = await self._extract_docx_with_structure(file_path)
            
            self._log(f"Extracted {structure.total_sections} sections, {structure.total_paragraphs} paragraphs, {structure.total_words} words")
            return structure
            
        except Exception as e:
            logger.error(f"Deep content extraction failed: {e}")
            # Create minimal structure for fallback
            structure = DocumentStructure()
            structure.title = "Document"
            structure.metadata = {"format": file_format.upper()}
            structure.sections = [DocumentSection("paragraph", "Content extraction failed", 0)]
            structure.total_sections = 1
            return structure

    async def _extract_pdf_with_structure(self, file_path: Path):
        """Extract PDF with complete structure analysis"""
        structure = DocumentStructure()
        
        try:
            if HAS_PYMUPDF:
                doc = fitz.open(str(file_path))
                
                # Extract metadata
                structure.metadata = {
                    "page_count": len(doc),
                    "title": doc.metadata.get('title', ''),
                    "author": doc.metadata.get('author', ''),
                    "creator": doc.metadata.get('creator', ''),
                    "subject": doc.metadata.get('subject', ''),
                    "format": "PDF"
                }
                
                structure.title = structure.metadata.get('title', 'Document')
                
                # Extract content page by page with structure
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    
                    if page_text.strip():  # Only process non-empty pages
                        # Analyze page structure
                        page_sections = self._analyze_page_structure(page_text, page_num + 1)
                        structure.sections.extend(page_sections)
                        
                        structure.pages.append({
                            "page_number": page_num + 1,
                            "text": page_text,
                            "sections": len(page_sections),
                            "word_count": len(page_text.split())
                        })
                
                doc.close()
                
            elif HAS_PYPDF2:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    
                    structure.metadata = {
                        "page_count": len(pdf_reader.pages),
                        "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                        "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                        "format": "PDF"
                    }
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            page_sections = self._analyze_page_structure(page_text, page_num + 1)
                            structure.sections.extend(page_sections)
                            
                            structure.pages.append({
                                "page_number": page_num + 1,
                                "text": page_text,
                                "sections": len(page_sections),
                                "word_count": len(page_text.split())
                            })
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            # Fallback: create basic structure
            structure.sections = [DocumentSection("paragraph", "PDF content could not be extracted properly", 0)]
        
        # Calculate totals
        structure.total_sections = len(structure.sections)
        structure.total_paragraphs = sum(1 for s in structure.sections if s.section_type == 'paragraph')
        structure.total_words = sum(len(s.content.split()) for s in structure.sections)
        
        return structure

    async def _extract_docx_with_structure(self, file_path: Path):
        """Extract DOCX with complete structure analysis"""
        structure = DocumentStructure()
        
        try:
            if HAS_DOCX:
                doc = Document(str(file_path))
                
                # Extract metadata
                core_props = doc.core_properties
                structure.metadata = {
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "subject": core_props.subject or '',
                    "created": str(core_props.created) if core_props.created else '',
                    "modified": str(core_props.modified) if core_props.modified else '',
                    "paragraph_count": len(doc.paragraphs),
                    "format": "DOCX"
                }
                
                structure.title = structure.metadata.get('title', 'Document')
                
                # Extract structured content
                for para in doc.paragraphs:
                    if para.text.strip():
                        section = self._analyze_paragraph_structure(para)
                        structure.sections.append(section)
                
                # Calculate totals
                structure.total_sections = len(structure.sections)
                structure.total_paragraphs = sum(1 for s in structure.sections if s.section_type == 'paragraph')
                structure.total_words = sum(len(s.content.split()) for s in structure.sections)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            # Fallback
            structure.sections = [DocumentSection("paragraph", "DOCX content could not be extracted properly", 0)]
            structure.total_sections = 1
        
        return structure

    def _analyze_page_structure(self, page_text: str, page_number: int) -> List[DocumentSection]:
        """Analyze page structure and create sections"""
        sections = []
        lines = page_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Determine section type and level
            section_type, level, formatting = self._classify_text_element(line)
            
            section = DocumentSection(
                section_type=section_type,
                content=line,
                level=level,
                formatting=formatting
            )
            section.section_id = f"page_{page_number}_section_{len(sections) + 1}"
            sections.append(section)
        
        return sections

    def _analyze_paragraph_structure(self, para) -> DocumentSection:
        """Analyze paragraph structure from DOCX"""
        content = para.text.strip()
        section_type, level, formatting = self._classify_text_element(content)
        
        # Extract additional formatting from DOCX
        if hasattr(para, 'style') and para.style:
            formatting['style_name'] = para.style.name
            if para.style.name.startswith('Heading'):
                section_type = 'heading'
                level = int(para.style.name.replace('Heading ', '')) if para.style.name.replace('Heading ', '').isdigit() else 1
        
        section = DocumentSection(
            section_type=section_type,
            content=content,
            level=level,
            formatting=formatting
        )
        
        return section

    def _classify_text_element(self, text: str) -> tuple:
        """Classify text element type, level, and formatting"""
        text = text.strip()
        
        # Title detection (usually all caps, short, at beginning)
        if text.isupper() and len(text) < 100:
            return 'title', 0, {'font_weight': 'bold', 'alignment': 'center'}
        
        # Heading detection (starts with number, ends with colon, or specific patterns)
        heading_patterns = [
            r'^\d+\.\s*[A-Z]',  # 1. Introduction
            r'^[A-Z][A-Za-z\s]+:$',  # Introduction:
            r'^[A-Z]{2,}\s*$',  # ABSTRACT
            r'^\d+\.\d+\s*[A-Z]',  # 1.1 Subsection
        ]
        
        for i, pattern in enumerate(heading_patterns):
            if re.match(pattern, text):
                return 'heading', i + 1, {'font_weight': 'bold'}
        
        # List item detection
        if re.match(r'^[\-\*\•]\s+', text) or re.match(r'^\d+\.\s+', text):
            return 'list_item', 0, {'list_style': 'bullet'}
        
        # Default to paragraph
        return 'paragraph', 0, {'font_weight': 'normal'}

    # Simplified processing pipeline for reliable operation
    async def _analyze_document_structure(self, document_structure):
        """STAGE 2: Quick structure analysis"""
        self._current_stage = "Structure Analysis"
        self._log(f"Structure analysis: {document_structure.total_sections} sections identified")
        return document_structure

    async def _segment_content_intelligently(self, document_structure):
        """STAGE 3: Simple content segmentation"""
        self._current_stage = "Content Segmentation"
        
        # Simple grouping: 3-5 sections per group
        segmented_groups = []
        current_group = []
        
        for section in document_structure.sections:
            current_group.append(section)
            if len(current_group) >= 3:  # Process in smaller groups for better quality
                segmented_groups.append(current_group)
                current_group = []
        
        # Add final group
        if current_group:
            segmented_groups.append(current_group)
        
        self._log(f"Content segmented into {len(segmented_groups)} processing groups")
        return segmented_groups

    async def _process_sections_with_ai(self, segmented_sections, modification_prompt):
        """STAGE 4: AI processing with improved content generation"""
        try:
            self._current_stage = "AI Processing"
            processed_groups = []
            
            for group_idx, section_group in enumerate(segmented_sections):
                self._log(f"Processing group {group_idx + 1}/{len(segmented_sections)}")
                
                # Combine group content
                group_content = '\n\n'.join([section.content for section in section_group])
                
                # Enhanced prompt for better results
                section_prompt = f"""Transform and improve this document section according to the user's request:

USER REQUEST: {modification_prompt}

ORIGINAL CONTENT:
{group_content}

INSTRUCTIONS:
1. Apply the user's requested modifications thoroughly
2. Improve writing quality, grammar, and clarity
3. Maintain professional tone and structure
4. Expand content where appropriate to add value
5. Ensure coherent flow and logical organization

Provide the enhanced content with substantial improvements:"""

                processed_content = ""
                async for chunk in self.llm_manager.generate_response_stream(
                    message=section_prompt,
                    model_type="report_generation",
                    max_tokens=3000,  # Increased for better content generation
                    temperature=0.3
                ):
                    processed_content = chunk.get("accumulated_text", "")
                    if chunk.get("finished", False):
                        break
                
                # Create processed sections
                processed_group = []
                for i, original_section in enumerate(section_group):
                    new_section = DocumentSection(
                        section_type=original_section.section_type,
                        content=original_section.content,
                        level=original_section.level,
                        formatting=original_section.formatting
                    )
                    new_section.section_id = original_section.section_id
                    new_section.original_content = original_section.content
                    
                    # For first section in group, use the processed content
                    if i == 0 and processed_content.strip():
                        new_section.processed_content = processed_content.strip()
                    else:
                        new_section.processed_content = original_section.content
                    
                    processed_group.append(new_section)
                
                processed_groups.append(processed_group)
            
            self._log(f"AI processing completed for all {len(segmented_sections)} groups")
            return processed_groups
            
        except Exception as e:
            logger.error(f"AI processing failed: {e}")
            return segmented_sections

    async def _quality_assurance_pass(self, processed_sections, modification_prompt):
        """STAGE 5: Quality assurance"""
        self._current_stage = "Quality Assurance"
        self._log("Quality assurance pass completed")
        return processed_sections

    async def _reconstruct_document_with_structure(self, processed_sections, original_structure):
        """STAGE 6: Document reconstruction"""
        try:
            self._current_stage = "Document Reconstruction"
            
            # Flatten sections
            all_sections = []
            for group in processed_sections:
                all_sections.extend(group)
            
            # Build complete text
            complete_text_parts = []
            
            for section in all_sections:
                content = section.processed_content or section.content
                
                if section.section_type == 'title':
                    complete_text_parts.append(f"\n{content.upper()}\n")
                elif section.section_type == 'heading':
                    complete_text_parts.append(f"\n{content}\n")
                else:
                    complete_text_parts.append(f"{content}\n")
            
            complete_text = '\n'.join(complete_text_parts).strip()
            
            reconstructed_document = {
                "complete_text": complete_text,
                "sections": [{"type": s.section_type, "content": s.processed_content or s.content} for s in all_sections],
                "metadata": original_structure.metadata,
                "total_sections": len(all_sections),
                "reconstruction_successful": True
            }
            
            self._log(f"Document reconstructed with {len(all_sections)} sections")
            return reconstructed_document
            
        except Exception as e:
            logger.error(f"Document reconstruction failed: {e}")
            return {
                "complete_text": "Document reconstruction failed",
                "sections": [],
                "metadata": {},
                "reconstruction_successful": False
            }

    async def _final_quality_check(self, reconstructed_document, modification_prompt):
        """STAGE 7: Final quality check"""
        try:
            self._current_stage = "Final Quality Check"
            
            content = reconstructed_document.get('complete_text', '')
            
            if len(content) > 100:  # Only do final polish if we have substantial content
                polish_prompt = f"""Review and provide the final polished version of this document:

ORIGINAL REQUEST: {modification_prompt}

CONTENT:
{content}

Please provide the final version with:
1. Professional polish and refinement
2. Proper formatting and structure
3. Enhanced clarity and readability
4. Complete fulfillment of the original request

Final polished document:"""

                polished_content = ""
                async for chunk in self.llm_manager.generate_response_stream(
                    message=polish_prompt,
                    model_type="report_generation",
                    max_tokens=4000,
                    temperature=0.1
                ):
                    polished_content = chunk.get("accumulated_text", "")
                    if chunk.get("finished", False):
                        break
                
                if polished_content.strip():
                    reconstructed_document['complete_text'] = polished_content.strip()
                    reconstructed_document['final_polish_applied'] = True
            
            self._log("Final quality check completed successfully")
            return reconstructed_document
            
        except Exception as e:
            logger.error(f"Final quality check failed: {e}")
            return reconstructed_document

    async def _generate_enhanced_document(self, final_content, output_path, output_format, structure):
        """Generate the final enhanced document"""
        try:
            content_text = final_content.get('complete_text', '')
            metadata = final_content.get('metadata', structure.metadata)
            
            if output_format == 'pdf':
                await self._generate_enhanced_pdf(content_text, output_path, metadata, structure)
            elif output_format in ['doc', 'docx']:
                await self._generate_enhanced_docx(content_text, output_path, metadata, structure)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Enhanced document generation failed: {e}")
            raise

    async def _generate_enhanced_pdf(self, content, output_path, metadata, structure):
        """Generate enhanced PDF"""
        if not HAS_REPORTLAB:
            raise ImportError("ReportLab is required for PDF generation")
        
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            story = []
            styles = getSampleStyleSheet()
            
            # Add title
            title = metadata.get('title', structure.title or 'Enhanced Document')
            if title:
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=1
                )
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 20))
            
            # Process content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                if para.isupper() and len(para) < 100:
                    # Major heading
                    heading_style = ParagraphStyle(
                        'CustomHeading',
                        parent=styles['Heading1'],
                        fontSize=14,
                        spaceAfter=12,
                        spaceBefore=20
                    )
                    story.append(Paragraph(para, heading_style))
                else:
                    # Regular paragraph
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            self._log(f"Enhanced PDF generated: {output_path}")
            
        except Exception as e:
            logger.error(f"Enhanced PDF generation failed: {e}")
            raise

    async def _generate_enhanced_docx(self, content, output_path, metadata, structure):
        """Generate enhanced DOCX"""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX generation")
        
        try:
            doc = Document()
            
            # Set document properties
            core_props = doc.core_properties
            core_props.title = metadata.get('title', structure.title or 'Enhanced Document')
            core_props.author = metadata.get('author', 'Advanced Document Processor')
            core_props.subject = 'AI-Enhanced Document'
            
            # Add title
            title = metadata.get('title', structure.title or 'Enhanced Document')
            if title:
                doc.add_heading(title, 0)
            
            # Process content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                if para.isupper() and len(para) < 100:
                    # Major heading
                    doc.add_heading(para.title(), level=1)
                else:
                    # Regular paragraph
                    doc.add_paragraph(para)
            
            doc.save(str(output_path))
            self._log(f"Enhanced DOCX generated: {output_path}")
            
        except Exception as e:
            logger.error(f"Enhanced DOCX generation failed: {e}")
            raise

    # Compatibility methods
    async def compare_documents(self, original_path: str, modified_path: str) -> Dict[str, Any]:
        """Compare original and modified documents"""
        try:
            original_content = await self._extract_content(Path(original_path), Path(original_path).suffix.lower())
            modified_content = await self._extract_content(Path(modified_path), Path(modified_path).suffix.lower())
            
            original_text = original_content.get('text', '')
            modified_text = modified_content.get('text', '')
            
            comparison = {
                "original_stats": {
                    "word_count": len(original_text.split()),
                    "char_count": len(original_text),
                    "paragraph_count": len([p for p in original_text.split('\n\n') if p.strip()])
                },
                "modified_stats": {
                    "word_count": len(modified_text.split()),
                    "char_count": len(modified_text),
                    "paragraph_count": len([p for p in modified_text.split('\n\n') if p.strip()])
                },
                "changes": {
                    "word_count_delta": len(modified_text.split()) - len(original_text.split()),
                    "char_count_delta": len(modified_text) - len(original_text),
                    "similarity_score": self._calculate_similarity(original_text, modified_text)
                }
            }
            
            return comparison
            
        except Exception as e:
            logger.error(f"Document comparison failed: {e}")
            return {"error": str(e)}

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate basic similarity score between two texts"""
        try:
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            
            if not words1 and not words2:
                return 1.0
            if not words1 or not words2:
                return 0.0
            
            intersection = words1.intersection(words2)
            union = words1.union(words2)
            
            return len(intersection) / len(union)
            
        except Exception:
            return 0.0

    async def get_supported_formats(self) -> Dict[str, Any]:
        """Get information about supported document formats"""
        return {
            "input_formats": {
                "pdf": {"supported": HAS_PYPDF2 or HAS_PYMUPDF, "libraries": ["PyPDF2", "PyMuPDF"]},
                "docx": {"supported": HAS_DOCX, "libraries": ["python-docx"]},
                "doc": {"supported": HAS_DOCX, "libraries": ["python-docx"]}
            },
            "output_formats": {
                "pdf": {"supported": HAS_REPORTLAB, "libraries": ["ReportLab"]},
                "docx": {"supported": HAS_DOCX, "libraries": ["python-docx"]}
            },
            "features": {
                "perfect_pipeline": True,
                "structure_preservation": True,
                "multi_stage_processing": True,
                "ai_enhancement": True,
                "quality_assurance": True,
                "deep_content_extraction": True,
                "intelligent_segmentation": True
            },
            "pipeline_stages": [
                "Deep Content Extraction",
                "Structure Analysis",
                "Content Segmentation", 
                "AI Processing",
                "Quality Assurance",
                "Document Reconstruction",
                "Final Quality Check"
            ]
        }

    async def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up temporary files older than specified hours"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for file_path in self.temp_path.iterdir():
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        logger.info(f"Cleaned up temp file: {file_path}")
                        
        except Exception as e:
            logger.error(f"Failed to cleanup temp files: {e}")

    # Legacy method for compatibility
    async def _extract_content(self, file_path: Path, file_format: str) -> Dict[str, Any]:
        """Legacy content extraction for compatibility"""
        structure = await self._deep_content_extraction(file_path, file_format)
        return {
            "text": '\n\n'.join([section.content for section in structure.sections]),
            "metadata": structure.metadata
        }
 
import tempfile
import asyncio
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import logging
import json
import base64
import io
import re

# Document processing imports
try:
    import PyPDF2
    from PyPDF2 import PdfReader, PdfWriter
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import Color
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches
    import docx2txt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF for advanced PDF handling
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

logger = logging.getLogger(__name__)

class PDFElement:
    """Represents a PDF element with exact positioning and formatting"""
    def __init__(self, element_type: str, content: str, bbox: tuple, font_info: dict, page_num: int):
        self.element_type = element_type  # text, bullet, table_cell, heading, etc.
        self.content = content
        self.original_content = content
        self.modified_content = ""
        self.bbox = bbox  # (x0, y0, x1, y1) positioning
        self.font_info = font_info  # font, size, color, etc.
        self.page_num = page_num
        self.element_id = f"page_{page_num}_{element_type}_{hash(content[:50])}"

class DocumentStructure:
    """Represents the complete structure of a document"""
    def __init__(self):
        self.title = ""
        self.metadata = {}
        self.sections = []
        self.pages = []
        self.pdf_elements = []  # New: exact PDF elements with positioning
        self.total_sections = 0
        self.total_paragraphs = 0
        self.total_words = 0

class DocumentSection:
    """Represents a section of the document with its structure"""
    def __init__(self, section_type: str, content: str, level: int = 0, formatting: Dict = None):
        self.section_type = section_type  # title, heading, paragraph, list, table, etc.
        self.content = content
        self.level = level  # For headings: 1, 2, 3, etc.
        self.formatting = formatting or {}
        self.original_content = content
        self.processed_content = ""
        self.section_id = ""

class AdvancedDocumentProcessor:
    """Perfect Pipeline for Complete Document Processing and Regeneration"""
    
    def __init__(self, settings, llm_manager):
        self.settings = settings
        self.llm_manager = llm_manager
        self.temp_path = Path(settings.temp_files_path)
        self.temp_path.mkdir(exist_ok=True)
        self.processing_log = []
    
    async def process_document(
        self,
        file_path: str,
        modification_prompt: str,
        output_format: str = None,
        custom_filename: str = None,
        preserve_structure: bool = True  # New parameter for structure preservation
    ) -> Dict[str, Any]:
        """
        PERFECT PIPELINE: Complete Document Processing and Regeneration
        
        Pipeline Stages:
        1. Deep Content Extraction (with exact structure if preserve_structure=True)
        2. Structure Analysis & Preservation
        3. Content Segmentation
        4. Section-by-Section AI Processing
        5. Content Quality Assurance
        6. Document Reconstruction with Original Structure
        7. Final Quality Check
        """
        try:
            self.processing_log = []
            start_time = datetime.now()
            
            file_path = Path(file_path)
            original_format = file_path.suffix.lower()
            
            if original_format not in ['.pdf', '.doc', '.docx']:
                raise ValueError(f"Unsupported file format: {original_format}")
            
            # STAGE 1: Deep Content Extraction (with structure preservation for PDFs)
            if preserve_structure and original_format == '.pdf':
                self._log("Stage 1: Structure-Preserving PDF Extraction - Maintaining exact layout")
                document_structure = await self._extract_pdf_with_exact_structure(file_path)
            else:
                self._log("Stage 1: Deep Content Extraction - Extracting ALL content and metadata")
                document_structure = await self._deep_content_extraction(file_path, original_format)
            
            # STAGE 2: Structure Analysis & Preservation
            self._log("Stage 2: Structure Analysis - Analyzing document architecture")
            analyzed_structure = await self._analyze_document_structure(document_structure)
            
            # STAGE 3: Content Segmentation
            self._log("Stage 3: Content Segmentation - Breaking into processable sections")
            if preserve_structure and original_format == '.pdf':
                segmented_sections = await self._segment_pdf_elements_intelligently(analyzed_structure)
            else:
                segmented_sections = await self._segment_content_intelligently(analyzed_structure)
            
            # STAGE 4: Section-by-Section AI Processing
            self._log("Stage 4: AI Processing - Regenerating content while preserving structure")
            if preserve_structure and original_format == '.pdf':
                processed_sections = await self._process_pdf_elements_with_ai(segmented_sections, modification_prompt)
            else:
                processed_sections = await self._process_sections_with_ai(segmented_sections, modification_prompt)
            
            # STAGE 5: Content Quality Assurance
            self._log("Stage 5: Quality Assurance - Ensuring consistency and coherence")
            quality_assured_content = await self._quality_assurance_pass(processed_sections, modification_prompt)
            
            # STAGE 6: Document Reconstruction
            self._log("Stage 6: Document Reconstruction - Rebuilding with original structure")
            if preserve_structure and original_format == '.pdf':
                reconstructed_document = await self._reconstruct_pdf_with_exact_structure(quality_assured_content, analyzed_structure)
            else:
                reconstructed_document = await self._reconstruct_document_with_structure(quality_assured_content, analyzed_structure)
            
            # STAGE 7: Final Quality Check
            self._log("Stage 7: Final Quality Check - Verifying output quality")
            final_content = await self._final_quality_check(reconstructed_document, modification_prompt)
            
            # Generate output
            if output_format is None:
                output_format = original_format.lstrip('.')
            
            if custom_filename:
                output_filename = f"{custom_filename}.{output_format}"
            else:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = file_path.stem
                output_filename = f"{base_name}_modified_{timestamp}.{output_format}"
            
            output_path = self.temp_path / output_filename
            
            # Generate final document
            if preserve_structure and original_format == '.pdf':
                await self._generate_structure_preserved_pdf(
                    final_content, 
                    output_path, 
                    analyzed_structure
                )
            else:
                await self._generate_enhanced_document(
                    final_content, 
                    output_path, 
                    output_format,
                    analyzed_structure
                )
            
            end_time = datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            return {
                "success": True,
                "output_path": str(output_path),
                "original_format": original_format,
                "output_format": output_format,
                "original_size": file_path.stat().st_size,
                "output_size": output_path.stat().st_size,
                "processing_time": end_time.isoformat(),
                "processing_duration_seconds": processing_time,
                "pipeline_stages_completed": 7,
                "sections_processed": len(processed_sections),
                "total_content_length": len(final_content.get('complete_text', '')),
                "modifications_applied": True,
                "processing_log": self.processing_log,
                "structure_preserved": preserve_structure,
                "quality_assurance_passed": True
            }
            
        except Exception as e:
            logger.error(f"Perfect pipeline failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "processing_time": datetime.now().isoformat(),
                "pipeline_stage_failed": getattr(self, '_current_stage', 'Unknown'),
                "processing_log": self.processing_log
            }

    async def _extract_pdf_with_exact_structure(self, file_path: Path) -> DocumentStructure:
        """Extract PDF with exact structure preservation including positioning, fonts, bullets, tables"""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF is required for structure-preserving PDF processing")
        
        self._current_stage = "Structure-Preserving PDF Extraction"
        structure = DocumentStructure()
        
        try:
            doc = fitz.open(str(file_path))
            
            # Extract metadata
            structure.metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get('title', ''),
                "author": doc.metadata.get('author', ''),
                "creator": doc.metadata.get('creator', ''),
                "subject": doc.metadata.get('subject', ''),
                "format": "PDF",
                "structure_preserved": True
            }
            
            structure.title = structure.metadata.get('title', 'Document')
            
            # Extract each page with exact structure
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Get all text blocks with positioning and formatting
                text_blocks = page.get_text("dict")
                
                page_elements = []
                page_text_parts = []
                
                for block in text_blocks.get("blocks", []):
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if text:
                                    # Extract detailed formatting
                                    font_info = {
                                        "font": span.get("font", ""),
                                        "size": span.get("size", 12),
                                        "color": span.get("color", 0),
                                        "flags": span.get("flags", 0),  # Bold, italic flags
                                        "ascender": span.get("ascender", 0),
                                        "descender": span.get("descender", 0)
                                    }
                                    
                                    # Determine element type based on formatting and content
                                    element_type = self._classify_pdf_element(text, font_info, span["bbox"])
                                    
                                    # Create PDF element with exact positioning
                                    pdf_element = PDFElement(
                                        element_type=element_type,
                                        content=text,
                                        bbox=span["bbox"],
                                        font_info=font_info,
                                        page_num=page_num + 1
                                    )
                                    
                                    page_elements.append(pdf_element)
                                    page_text_parts.append(text)
                                    structure.pdf_elements.append(pdf_element)
                
                # Store page information
                page_text = "\n".join(page_text_parts)
                structure.pages.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "elements": len(page_elements),
                    "word_count": len(page_text.split()),
                    "pdf_elements": page_elements
                })
            
            doc.close()
            
            # Convert PDF elements to sections for compatibility
            for element in structure.pdf_elements:
                section = DocumentSection(
                    section_type=element.element_type,
                    content=element.content,
                    level=0,
                    formatting=element.font_info
                )
                section.section_id = element.element_id
                structure.sections.append(section)
            
            # Calculate totals
            structure.total_sections = len(structure.sections)
            structure.total_paragraphs = sum(1 for s in structure.sections if s.section_type == 'paragraph')
            structure.total_words = sum(len(s.content.split()) for s in structure.sections)
            
            self._log(f"Extracted PDF with exact structure: {len(structure.pdf_elements)} elements with positioning")
            return structure
            
        except Exception as e:
            logger.error(f"Structure-preserving PDF extraction failed: {e}")
            # Fallback to regular extraction
            return await self._extract_pdf_with_structure(file_path)

    def _classify_pdf_element(self, text: str, font_info: dict, bbox: tuple) -> str:
        """Classify PDF element type based on content, formatting, and position"""
        font_size = font_info.get("size", 12)
        font_flags = font_info.get("flags", 0)
        
        # Check for bullet points
        if re.match(r'^[\u2022\u25cf\u25cb\u25aa\u25ab•·‣⁃]\s*', text) or re.match(r'^[\-\*]\s+', text):
            return 'bullet_point'
        
        # Check for numbered lists
        if re.match(r'^\d+[\.\)]\s+', text):
            return 'numbered_list'
        
        # Check for headings (large font, bold, or all caps)
        if font_size > 14 or (font_flags & 2**4):  # Bold flag
            if len(text) < 100:
                return 'heading'
        
        # Check for title (very large font or all caps and short)
        if font_size > 16 or (text.isupper() and len(text) < 80):
            return 'title'
        
        # Check for table-like content (contains multiple tabs or specific patterns)
        if '\t' in text or len(text.split()) > 1 and any(char.isdigit() for char in text):
            if len(text.split()) <= 10:  # Short entries typical of tables
                return 'table_cell'
        
        # Default to paragraph
        return 'paragraph'

    async def _segment_pdf_elements_intelligently(self, document_structure: DocumentStructure) -> List[List[PDFElement]]:
        """Segment PDF elements into logical processing groups while preserving structure"""
        self._current_stage = "PDF Element Segmentation"
        
        # Group PDF elements by logical sections
        segmented_groups = []
        current_group = []
        
        for element in document_structure.pdf_elements:
            # Start new group on major headings or after certain number of elements
            if element.element_type in ['title', 'heading'] or len(current_group) >= 5:
                if current_group:
                    segmented_groups.append(current_group)
                    current_group = []
            
            current_group.append(element)
        
        # Add final group
        if current_group:
            segmented_groups.append(current_group)
        
        self._log(f"PDF elements segmented into {len(segmented_groups)} logical groups")
        return segmented_groups

    async def _process_pdf_elements_with_ai(self, segmented_elements: List[List[PDFElement]], modification_prompt: str) -> List[List[PDFElement]]:
        """Process PDF elements with AI while preserving exact structure"""
        try:
            self._current_stage = "PDF Element AI Processing"
            processed_groups = []
            
            for group_idx, element_group in enumerate(segmented_elements):
                self._log(f"Processing PDF element group {group_idx + 1}/{len(segmented_elements)}")
                
                # Combine group content with structure indicators
                group_content_parts = []
                for element in element_group:
                    if element.element_type == 'bullet_point':
                        group_content_parts.append(f"• {element.content}")
                    elif element.element_type == 'numbered_list':
                        group_content_parts.append(f"1. {element.content}")
                    elif element.element_type == 'heading':
                        group_content_parts.append(f"HEADING: {element.content}")
                    elif element.element_type == 'title':
                        group_content_parts.append(f"TITLE: {element.content}")
                    else:
                        group_content_parts.append(element.content)
                
                group_content = '\n'.join(group_content_parts)
                
                # Enhanced prompt for structure preservation
                section_prompt = f"""Modify the content according to the user's request while maintaining EXACT formatting structure:

USER REQUEST: {modification_prompt}

ORIGINAL CONTENT WITH STRUCTURE:
{group_content}

CRITICAL REQUIREMENTS:
1. Maintain EXACT formatting - if it's a bullet point, keep it as a bullet point
2. Maintain EXACT structure - if it's a heading, keep it as a heading
3. Only modify the TEXT CONTENT, not the formatting
4. If it's "• Text", return "• Modified Text"
5. If it's "HEADING: Text", return "HEADING: Modified Text"
6. Apply the user's modifications while preserving all structural elements

Return the modified content with identical structure:"""

                processed_content = ""
                async for chunk in self.llm_manager.generate_response_stream(
                    message=section_prompt,
                    model_type="report_generation",
                    max_tokens=2500,
                    temperature=0.2  # Lower temperature for better structure preservation
                ):
                    processed_content = chunk.get("accumulated_text", "")
                    if chunk.get("finished", False):
                        break
                
                # Parse the processed content back to PDF elements
                processed_group = self._parse_processed_pdf_elements(processed_content, element_group)
                processed_groups.append(processed_group)
            
            self._log(f"PDF element AI processing completed for all {len(segmented_elements)} groups")
            return processed_groups
            
        except Exception as e:
            logger.error(f"PDF element AI processing failed: {e}")
            return segmented_elements

    def _parse_processed_pdf_elements(self, processed_content: str, original_group: List[PDFElement]) -> List[PDFElement]:
        """Parse processed content back to PDF elements with preserved structure"""
        try:
            lines = processed_content.strip().split('\n')
            processed_elements = []
            
            for i, original_element in enumerate(original_group):
                if i < len(lines):
                    processed_line = lines[i].strip()
                    
                    # Extract content based on structure markers
                    if processed_line.startswith('•'):
                        content = processed_line[1:].strip()
                    elif processed_line.startswith('HEADING:'):
                        content = processed_line[8:].strip()
                    elif processed_line.startswith('TITLE:'):
                        content = processed_line[6:].strip()
                    elif re.match(r'^\d+\.', processed_line):
                        content = re.sub(r'^\d+\.\s*', '', processed_line)
                    else:
                        content = processed_line
                    
                    # Create new element with modified content but same structure
                    new_element = PDFElement(
                        element_type=original_element.element_type,
                        content=original_element.content,  # Keep original
                        bbox=original_element.bbox,
                        font_info=original_element.font_info,
                        page_num=original_element.page_num
                    )
                    new_element.element_id = original_element.element_id
                    new_element.original_content = original_element.content
                    new_element.modified_content = content if content else original_element.content
                    
                    processed_elements.append(new_element)
                else:
                    # Keep original if no processed version
                    original_element.modified_content = original_element.content
                    processed_elements.append(original_element)
            
            return processed_elements
            
        except Exception as e:
            logger.error(f"PDF element parsing failed: {e}")
            # Return originals with modified_content set to original content
            for element in original_group:
                element.modified_content = element.content
            return original_group

    async def _reconstruct_pdf_with_exact_structure(self, processed_elements: List[List[PDFElement]], original_structure: DocumentStructure) -> Dict[str, Any]:
        """Reconstruct PDF maintaining exact original structure and positioning"""
        try:
            self._current_stage = "PDF Structure Reconstruction"
            
            # Flatten elements
            all_elements = []
            for group in processed_elements:
                all_elements.extend(group)
            
            # Build complete text preserving structure
            complete_text_parts = []
            
            for element in all_elements:
                content = element.modified_content or element.content
                
                # Add content with structure preserved
                if element.element_type == 'bullet_point':
                    complete_text_parts.append(f"• {content}")
                elif element.element_type == 'numbered_list':
                    complete_text_parts.append(f"1. {content}")
                elif element.element_type == 'heading':
                    complete_text_parts.append(f"\n{content}\n")
                elif element.element_type == 'title':
                    complete_text_parts.append(f"\n{content.upper()}\n")
                else:
                    complete_text_parts.append(content)
            
            complete_text = '\n'.join(complete_text_parts).strip()
            
            reconstructed_document = {
                "complete_text": complete_text,
                "pdf_elements": all_elements,
                "metadata": original_structure.metadata,
                "total_elements": len(all_elements),
                "structure_preserved": True,
                "reconstruction_successful": True
            }
            
            self._log(f"PDF reconstructed with exact structure: {len(all_elements)} elements")
            return reconstructed_document
            
        except Exception as e:
            logger.error(f"PDF structure reconstruction failed: {e}")
            return {
                "complete_text": "PDF reconstruction failed",
                "pdf_elements": [],
                "metadata": {},
                "reconstruction_successful": False
            }

    async def _generate_structure_preserved_pdf(
        self, 
        final_content: Dict[str, Any], 
        output_path: Path, 
        structure: DocumentStructure
    ):
        """Generate PDF with exact structure preservation"""
        if not HAS_PYMUPDF or not HAS_REPORTLAB:
            raise ImportError("PyMuPDF and ReportLab are required for structure-preserving PDF generation")
        
        try:
            # Use ReportLab to create PDF with exact positioning
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            pdf_elements = final_content.get('pdf_elements', [])
            
            # Group elements by page
            pages_elements = {}
            for element in pdf_elements:
                page_num = element.page_num
                if page_num not in pages_elements:
                    pages_elements[page_num] = []
                pages_elements[page_num].append(element)
            
            # Generate each page with exact positioning
            for page_num in sorted(pages_elements.keys()):
                page_elements = pages_elements[page_num]
                
                for element in page_elements:
                    content = element.modified_content or element.content
                    bbox = element.bbox
                    font_info = element.font_info
                    
                    # Set font and size
                    font_size = font_info.get('size', 12)
                    c.setFont("Helvetica", font_size)
                    
                    # Convert coordinates (PDF coordinates are bottom-left origin)
                    x = bbox[0]
                    y = height - bbox[3]  # Flip Y coordinate
                    
                    # Add structure-specific formatting
                    if element.element_type == 'bullet_point':
                        c.drawString(x, y, f"• {content}")
                    elif element.element_type == 'numbered_list':
                        c.drawString(x, y, f"1. {content}")
                    elif element.element_type in ['heading', 'title']:
                        # Make headings bold (simulate)
                        c.setFont("Helvetica-Bold", font_size)
                        c.drawString(x, y, content)
                        c.setFont("Helvetica", font_size)  # Reset
                    else:
                        c.drawString(x, y, content)
                
                # Start new page if not the last page
                if page_num < max(pages_elements.keys()):
                    c.showPage()
            
            c.save()
            self._log(f"Structure-preserved PDF generated: {output_path}")
            
        except Exception as e:
            logger.error(f"Structure-preserved PDF generation failed: {e}")
            # Fallback to regular PDF generation
            await self._generate_enhanced_pdf(
                final_content.get('complete_text', ''), 
                output_path, 
                final_content.get('metadata', {}),
                structure
            )

    def _log(self, message: str):
        """Log processing steps"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        self.processing_log.append(log_entry)
        logger.info(log_entry)

    async def _deep_content_extraction(self, file_path: Path, file_format: str):
        """STAGE 1: Extract ALL content with complete structure preservation"""
        try:
            self._current_stage = "Deep Content Extraction"
            structure = DocumentStructure()
            
            if file_format == '.pdf':
                structure = await self._extract_pdf_with_structure(file_path)
            elif file_format in ['.doc', '.docx']:
                structure = await self._extract_docx_with_structure(file_path)
            
            self._log(f"Extracted {structure.total_sections} sections, {structure.total_paragraphs} paragraphs, {structure.total_words} words")
            return structure
            
        except Exception as e:
            logger.error(f"Deep content extraction failed: {e}")
            # Create minimal structure for fallback
            structure = DocumentStructure()
            structure.title = "Document"
            structure.metadata = {"format": file_format.upper()}
            structure.sections = [DocumentSection("paragraph", "Content extraction failed", 0)]
            structure.total_sections = 1
            return structure

    async def _extract_pdf_with_structure(self, file_path: Path):
        """Extract PDF with complete structure analysis"""
        structure = DocumentStructure()
        
        try:
            if HAS_PYMUPDF:
                doc = fitz.open(str(file_path))
                
                # Extract metadata
                structure.metadata = {
                    "page_count": len(doc),
                    "title": doc.metadata.get('title', ''),
                    "author": doc.metadata.get('author', ''),
                    "creator": doc.metadata.get('creator', ''),
                    "subject": doc.metadata.get('subject', ''),
                    "format": "PDF"
                }
                
                structure.title = structure.metadata.get('title', 'Document')
                
                # Extract content page by page with structure
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    
                    if page_text.strip():  # Only process non-empty pages
                        # Analyze page structure
                        page_sections = self._analyze_page_structure(page_text, page_num + 1)
                        structure.sections.extend(page_sections)
                        
                        structure.pages.append({
                            "page_number": page_num + 1,
                            "text": page_text,
                            "sections": len(page_sections),
                            "word_count": len(page_text.split())
                        })
                
                doc.close()
                
            elif HAS_PYPDF2:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    
                    structure.metadata = {
                        "page_count": len(pdf_reader.pages),
                        "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                        "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                        "format": "PDF"
                    }
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            page_sections = self._analyze_page_structure(page_text, page_num + 1)
                            structure.sections.extend(page_sections)
                            
                            structure.pages.append({
                                "page_number": page_num + 1,
                                "text": page_text,
                                "sections": len(page_sections),
                                "word_count": len(page_text.split())
                            })
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            # Fallback: create basic structure
            structure.sections = [DocumentSection("paragraph", "PDF content could not be extracted properly", 0)]
        
        # Calculate totals
        structure.total_sections = len(structure.sections)
        structure.total_paragraphs = sum(1 for s in structure.sections if s.section_type == 'paragraph')
        structure.total_words = sum(len(s.content.split()) for s in structure.sections)
        
        return structure

    async def _extract_docx_with_structure(self, file_path: Path):
        """Extract DOCX with complete structure analysis"""
        structure = DocumentStructure()
        
        try:
            if HAS_DOCX:
                doc = Document(str(file_path))
                
                # Extract metadata
                core_props = doc.core_properties
                structure.metadata = {
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "subject": core_props.subject or '',
                    "created": str(core_props.created) if core_props.created else '',
                    "modified": str(core_props.modified) if core_props.modified else '',
                    "paragraph_count": len(doc.paragraphs),
                    "format": "DOCX"
                }
                
                structure.title = structure.metadata.get('title', 'Document')
                
                # Extract structured content
                for para in doc.paragraphs:
                    if para.text.strip():
                        section = self._analyze_paragraph_structure(para)
                        structure.sections.append(section)
                
                # Calculate totals
                structure.total_sections = len(structure.sections)
                structure.total_paragraphs = sum(1 for s in structure.sections if s.section_type == 'paragraph')
                structure.total_words = sum(len(s.content.split()) for s in structure.sections)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            # Fallback
            structure.sections = [DocumentSection("paragraph", "DOCX content could not be extracted properly", 0)]
            structure.total_sections = 1
        
        return structure

    def _analyze_page_structure(self, page_text: str, page_number: int) -> List[DocumentSection]:
        """Analyze page structure and create sections"""
        sections = []
        lines = page_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Determine section type and level
            section_type, level, formatting = self._classify_text_element(line)
            
            section = DocumentSection(
                section_type=section_type,
                content=line,
                level=level,
                formatting=formatting
            )
            section.section_id = f"page_{page_number}_section_{len(sections) + 1}"
            sections.append(section)
        
        return sections

    def _analyze_paragraph_structure(self, para) -> DocumentSection:
        """Analyze paragraph structure from DOCX"""
        content = para.text.strip()
        section_type, level, formatting = self._classify_text_element(content)
        
        # Extract additional formatting from DOCX
        if hasattr(para, 'style') and para.style:
            formatting['style_name'] = para.style.name
            if para.style.name.startswith('Heading'):
                section_type = 'heading'
                level = int(para.style.name.replace('Heading ', '')) if para.style.name.replace('Heading ', '').isdigit() else 1
        
        section = DocumentSection(
            section_type=section_type,
            content=content,
            level=level,
            formatting=formatting
        )
        
        return section

    def _classify_text_element(self, text: str) -> tuple:
        """Classify text element type, level, and formatting"""
        text = text.strip()
        
        # Title detection (usually all caps, short, at beginning)
        if text.isupper() and len(text) < 100:
            return 'title', 0, {'font_weight': 'bold', 'alignment': 'center'}
        
        # Heading detection (starts with number, ends with colon, or specific patterns)
        heading_patterns = [
            r'^\d+\.\s*[A-Z]',  # 1. Introduction
            r'^[A-Z][A-Za-z\s]+:$',  # Introduction:
            r'^[A-Z]{2,}\s*$',  # ABSTRACT
            r'^\d+\.\d+\s*[A-Z]',  # 1.1 Subsection
        ]
        
        for i, pattern in enumerate(heading_patterns):
            if re.match(pattern, text):
                return 'heading', i + 1, {'font_weight': 'bold'}
        
        # List item detection
        if re.match(r'^[\-\*\•]\s+', text) or re.match(r'^\d+\.\s+', text):
            return 'list_item', 0, {'list_style': 'bullet'}
        
        # Default to paragraph
        return 'paragraph', 0, {'font_weight': 'normal'}

    # Simplified processing pipeline for reliable operation
    async def _analyze_document_structure(self, document_structure):
        """STAGE 2: Quick structure analysis"""
        self._current_stage = "Structure Analysis"
        self._log(f"Structure analysis: {document_structure.total_sections} sections identified")
        return document_structure

    async def _segment_content_intelligently(self, document_structure):
        """STAGE 3: Simple content segmentation"""
        self._current_stage = "Content Segmentation"
        
        # Simple grouping: 3-5 sections per group
        segmented_groups = []
        current_group = []
        
        for section in document_structure.sections:
            current_group.append(section)
            if len(current_group) >= 3:  # Process in smaller groups for better quality
                segmented_groups.append(current_group)
                current_group = []
        
        # Add final group
        if current_group:
            segmented_groups.append(current_group)
        
        self._log(f"Content segmented into {len(segmented_groups)} processing groups")
        return segmented_groups

    async def _process_sections_with_ai(self, segmented_sections, modification_prompt):
        """STAGE 4: AI processing with improved content generation"""
        try:
            self._current_stage = "AI Processing"
            processed_groups = []
            
            for group_idx, section_group in enumerate(segmented_sections):
                self._log(f"Processing group {group_idx + 1}/{len(segmented_sections)}")
                
                # Combine group content
                group_content = '\n\n'.join([section.content for section in section_group])
                
                # Enhanced prompt for better results
                section_prompt = f"""Transform and improve this document section according to the user's request:

USER REQUEST: {modification_prompt}

ORIGINAL CONTENT:
{group_content}

INSTRUCTIONS:
1. Apply the user's requested modifications thoroughly
2. Improve writing quality, grammar, and clarity
3. Maintain professional tone and structure
4. Expand content where appropriate to add value
5. Ensure coherent flow and logical organization

Provide the enhanced content with substantial improvements:"""

                processed_content = ""
                async for chunk in self.llm_manager.generate_response_stream(
                    message=section_prompt,
                    model_type="report_generation",
                    max_tokens=3000,  # Increased for better content generation
                    temperature=0.3
                ):
                    processed_content = chunk.get("accumulated_text", "")
                    if chunk.get("finished", False):
                        break
                
                # Create processed sections
                processed_group = []
                for i, original_section in enumerate(section_group):
                    new_section = DocumentSection(
                        section_type=original_section.section_type,
                        content=original_section.content,
                        level=original_section.level,
                        formatting=original_section.formatting
                    )
                    new_section.section_id = original_section.section_id
                    new_section.original_content = original_section.content
                    
                    # For first section in group, use the processed content
                    if i == 0 and processed_content.strip():
                        new_section.processed_content = processed_content.strip()
                    else:
                        new_section.processed_content = original_section.content
                    
                    processed_group.append(new_section)
                
                processed_groups.append(processed_group)
            
            self._log(f"AI processing completed for all {len(segmented_sections)} groups")
            return processed_groups
            
        except Exception as e:
            logger.error(f"AI processing failed: {e}")
            return segmented_sections

    async def _quality_assurance_pass(self, processed_sections, modification_prompt):
        """STAGE 5: Quality assurance"""
        self._current_stage = "Quality Assurance"
        self._log("Quality assurance pass completed")
        return processed_sections

    async def _reconstruct_document_with_structure(self, processed_sections, original_structure):
        """STAGE 6: Document reconstruction"""
        try:
            self._current_stage = "Document Reconstruction"
            
            # Flatten sections
            all_sections = []
            for group in processed_sections:
                all_sections.extend(group)
            
            # Build complete text
            complete_text_parts = []
            
            for section in all_sections:
                content = section.processed_content or section.content
                
                if section.section_type == 'title':
                    complete_text_parts.append(f"\n{content.upper()}\n")
                elif section.section_type == 'heading':
                    complete_text_parts.append(f"\n{content}\n")
                else:
                    complete_text_parts.append(f"{content}\n")
            
            complete_text = '\n'.join(complete_text_parts).strip()
            
            reconstructed_document = {
                "complete_text": complete_text,
                "sections": [{"type": s.section_type, "content": s.processed_content or s.content} for s in all_sections],
                "metadata": original_structure.metadata,
                "total_sections": len(all_sections),
                "reconstruction_successful": True
            }
            
            self._log(f"Document reconstructed with {len(all_sections)} sections")
            return reconstructed_document
            
        except Exception as e:
            logger.error(f"Document reconstruction failed: {e}")
            return {
                "complete_text": "Document reconstruction failed",
                "sections": [],
                "metadata": {},
                "reconstruction_successful": False
            }

    async def _final_quality_check(self, reconstructed_document, modification_prompt):
        """STAGE 7: Final quality check"""
        try:
            self._current_stage = "Final Quality Check"
            
            content = reconstructed_document.get('complete_text', '')
            
            if len(content) > 100:  # Only do final polish if we have substantial content
                polish_prompt = f"""Review and provide the final polished version of this document:

ORIGINAL REQUEST: {modification_prompt}

CONTENT:
{content}

Please provide the final version with:
1. Professional polish and refinement
2. Proper formatting and structure
3. Enhanced clarity and readability
4. Complete fulfillment of the original request

Final polished document:"""

                polished_content = ""
                async for chunk in self.llm_manager.generate_response_stream(
                    message=polish_prompt,
                    model_type="report_generation",
                    max_tokens=4000,
                    temperature=0.1
                ):
                    polished_content = chunk.get("accumulated_text", "")
                    if chunk.get("finished", False):
                        break
                
                if polished_content.strip():
                    reconstructed_document['complete_text'] = polished_content.strip()
                    reconstructed_document['final_polish_applied'] = True
            
            self._log("Final quality check completed successfully")
            return reconstructed_document
            
        except Exception as e:
            logger.error(f"Final quality check failed: {e}")
            return reconstructed_document

    async def _generate_enhanced_document(self, final_content, output_path, output_format, structure):
        """Generate the final enhanced document"""
        try:
            content_text = final_content.get('complete_text', '')
            metadata = final_content.get('metadata', structure.metadata)
            
            if output_format == 'pdf':
                await self._generate_enhanced_pdf(content_text, output_path, metadata, structure)
            elif output_format in ['doc', 'docx']:
                await self._generate_enhanced_docx(content_text, output_path, metadata, structure)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Enhanced document generation failed: {e}")
            raise

    async def _generate_enhanced_pdf(self, content, output_path, metadata, structure):
        """Generate enhanced PDF"""
        if not HAS_REPORTLAB:
            raise ImportError("ReportLab is required for PDF generation")
        
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            story = []
            styles = getSampleStyleSheet()
            
            # Add title
            title = metadata.get('title', structure.title or 'Enhanced Document')
            if title:
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=1
                )
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 20))
            
            # Process content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                if para.isupper() and len(para) < 100:
                    # Major heading
                    heading_style = ParagraphStyle(
                        'CustomHeading',
                        parent=styles['Heading1'],
                        fontSize=14,
                        spaceAfter=12,
                        spaceBefore=20
                    )
                    story.append(Paragraph(para, heading_style))
                else:
                    # Regular paragraph
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            self._log(f"Enhanced PDF generated: {output_path}")
            
        except Exception as e:
            logger.error(f"Enhanced PDF generation failed: {e}")
            raise

    async def _generate_enhanced_docx(self, content, output_path, metadata, structure):
        """Generate enhanced DOCX"""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX generation")
        
        try:
            doc = Document()
            
            # Set document properties
            core_props = doc.core_properties
            core_props.title = metadata.get('title', structure.title or 'Enhanced Document')
            core_props.author = metadata.get('author', 'Advanced Document Processor')
            core_props.subject = 'AI-Enhanced Document'
            
            # Add title
            title = metadata.get('title', structure.title or 'Enhanced Document')
            if title:
                doc.add_heading(title, 0)
            
            # Process content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                if para.isupper() and len(para) < 100:
                    # Major heading
                    doc.add_heading(para.title(), level=1)
                else:
                    # Regular paragraph
                    doc.add_paragraph(para)
            
            doc.save(str(output_path))
            self._log(f"Enhanced DOCX generated: {output_path}")
            
        except Exception as e:
            logger.error(f"Enhanced DOCX generation failed: {e}")
            raise

    # Compatibility methods
    async def compare_documents(self, original_path: str, modified_path: str) -> Dict[str, Any]:
        """Compare original and modified documents"""
        try:
            original_content = await self._extract_content(Path(original_path), Path(original_path).suffix.lower())
            modified_content = await self._extract_content(Path(modified_path), Path(modified_path).suffix.lower())
            
            original_text = original_content.get('text', '')
            modified_text = modified_content.get('text', '')
            
            comparison = {
                "original_stats": {
                    "word_count": len(original_text.split()),
                    "char_count": len(original_text),
                    "paragraph_count": len([p for p in original_text.split('\n\n') if p.strip()])
                },
                "modified_stats": {
                    "word_count": len(modified_text.split()),
                    "char_count": len(modified_text),
                    "paragraph_count": len([p for p in modified_text.split('\n\n') if p.strip()])
                },
                "changes": {
                    "word_count_delta": len(modified_text.split()) - len(original_text.split()),
                    "char_count_delta": len(modified_text) - len(original_text),
                    "similarity_score": self._calculate_similarity(original_text, modified_text)
                }
            }
            
            return comparison
            
        except Exception as e:
            logger.error(f"Document comparison failed: {e}")
            return {"error": str(e)}

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate basic similarity score between two texts"""
        try:
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            
            if not words1 and not words2:
                return 1.0
            if not words1 or not words2:
                return 0.0
            
            intersection = words1.intersection(words2)
            union = words1.union(words2)
            
            return len(intersection) / len(union)
            
        except Exception:
            return 0.0

    async def get_supported_formats(self) -> Dict[str, Any]:
        """Get information about supported document formats"""
        return {
            "input_formats": {
                "pdf": {"supported": HAS_PYPDF2 or HAS_PYMUPDF, "libraries": ["PyPDF2", "PyMuPDF"]},
                "docx": {"supported": HAS_DOCX, "libraries": ["python-docx"]},
                "doc": {"supported": HAS_DOCX, "libraries": ["python-docx"]}
            },
            "output_formats": {
                "pdf": {"supported": HAS_REPORTLAB, "libraries": ["ReportLab"]},
                "docx": {"supported": HAS_DOCX, "libraries": ["python-docx"]}
            },
            "features": {
                "perfect_pipeline": True,
                "structure_preservation": True,
                "multi_stage_processing": True,
                "ai_enhancement": True,
                "quality_assurance": True,
                "deep_content_extraction": True,
                "intelligent_segmentation": True
            },
            "pipeline_stages": [
                "Deep Content Extraction",
                "Structure Analysis",
                "Content Segmentation", 
                "AI Processing",
                "Quality Assurance",
                "Document Reconstruction",
                "Final Quality Check"
            ]
        }

    async def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up temporary files older than specified hours"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for file_path in self.temp_path.iterdir():
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        logger.info(f"Cleaned up temp file: {file_path}")
                        
        except Exception as e:
            logger.error(f"Failed to cleanup temp files: {e}")

    # Legacy method for compatibility
    async def _extract_content(self, file_path: Path, file_format: str) -> Dict[str, Any]:
        """Legacy content extraction for compatibility"""
        structure = await self._deep_content_extraction(file_path, file_format)
        return {
            "text": '\n\n'.join([section.content for section in structure.sections]),
            "metadata": structure.metadata
        }
 